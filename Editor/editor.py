import json
import os
import re
import shutil
import sys
import uuid
import argparse
import urllib.parse
import base64
import threading
import hashlib
import hmac
import struct
import requests
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from typing import Dict, List, Optional, Tuple, Set, Callable, Any


# =====================
# AppData Logging
# =====================
def _get_appdata_dir() -> str:
    # Windows: %APPDATA%\storyweaver. Fallbacks for safety.
    base = os.environ.get("APPDATA")
    if not base:
        base = os.path.join(os.path.expanduser("~"), "AppData", "Roaming")
    d = os.path.join(base, "storyweaver")
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass
    return d

LOG_DIR = _get_appdata_dir()
LOG_PATH = os.path.join(LOG_DIR, "editor.log")

def log_event(event: str, payload: dict = None, **payload_kw):
    """
    AppData logger.

    Backward compatible with both:
      - log_event("EVT", a=1, b=2)
      - log_event("EVT", {"a": 1, "b": 2})
    """
    try:
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        merged = {}
        if isinstance(payload, dict):
            merged.update(payload)
        elif payload is not None:
            merged["payload"] = payload
        merged.update(payload_kw)

        rec = {"event": event, **merged}
        line = f"[{ts}] {event} {json.dumps(rec, ensure_ascii=False, default=str)}\n"
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line)
    except Exception:
        pass

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtWebSockets import QWebSocket
from PySide6.QtNetwork import QNetworkRequest

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
try:
    from docx import Document
except Exception:
    Document = None


# =========================
# Helpers / Normalization
# =========================

SCHEMA_VERSION = 3  # unchanged (v3 = goals)

APP_NAME = "StoryWeaver"
# Версия приложения (важно для лаунчера: editor/version.txt)
# ДОЛЖНА совпадать с manifest.json -> editor.version на сервере.
APP_VERSION = "3.1"

def get_app_version() -> str:
    return str(APP_VERSION)

def _get_install_dir() -> str:
    # В PyInstaller onefile/__file__ может указывать на временную папку распаковки.
    # Нам нужна директория, где лежит editor.exe (или editor.py в dev-режиме).
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))

ICON_PATH = os.path.join(_get_install_dir(), "storyweaver.ico")
if not os.path.exists(ICON_PATH):
    ICON_PATH = os.path.join(os.path.dirname(_get_install_dir()), "storyweaver.ico")

def write_version_txt():
    """Пишет editor/version.txt (рядом с editor.exe), чтобы лаунчер мог прочитать версию без запуска редактора."""
    try:
        d = _get_install_dir()
        path = os.path.join(d, "version.txt")
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            f.write(get_app_version().strip() + "\n")
        os.replace(tmp, path)
        log_event("VERSION_TXT_WRITTEN", path=path, version=get_app_version())
    except Exception as e:
        # Не критично: лаунчер умеет читать версию из ресурсов .exe, если будет.
        try:
            log_event("VERSION_TXT_WRITE_FAIL", err=str(e))
        except Exception:
            pass


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def new_uuid() -> str:
    return str(uuid.uuid4())

def norm_spaces(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s

def norm_key(s: str) -> str:
    return norm_spaces(s).casefold()

def words_count(s: str) -> int:
    return 0 if not s else len([w for w in s.strip().split(" ") if w])

LINK_RE = re.compile(r"::(.+?)::")

def extract_link_at_pos(text: str, pos: int) -> Optional[Tuple[int, int, str]]:
    for m in LINK_RE.finditer(text):
        s, e = m.span()
        if s <= pos <= e:
            inner = m.group(1)
            return s, e, inner
    return None

def safe_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[<>:"/\\|?*]+', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "Unnamed"
    return name

def strip_link_markers(text: str) -> str:
    """Remove technical ::Name:: markers for PDF and other outputs."""
    return LINK_RE.sub(lambda m: m.group(1), text or "")


# =========================
# Resource path (PyInstaller safe)
# =========================

def resource_path(rel_path: str) -> str:
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel_path)


# =========================
# Data Model
# =========================

@dataclass
class Meta:
    created_at: str = field(default_factory=now_iso)
    updated_at: str = field(default_factory=now_iso)
    updated_by: str = "local"
    revision: int = 1
    is_deleted: bool = False

    def touch(self, user: str = "local"):
        self.updated_at = now_iso()
        self.updated_by = user
        self.revision += 1


@dataclass
class CustomTagField:
    field_title: str = ""
    values: List[str] = field(default_factory=list)


@dataclass
class Visual:
    x: float = 0.0
    y: float = 0.0


@dataclass
class Character:
    id: str = field(default_factory=new_uuid)
    name: str = ""
    story_public: str = ""
    story_private: str = ""
    factions: List[str] = field(default_factory=list)
    locations: List[str] = field(default_factory=list)
    stories: List[str] = field(default_factory=list)
    masters: List[str] = field(default_factory=list)
    custom_tag_fields: List[CustomTagField] = field(default_factory=lambda: [CustomTagField() for _ in range(5)])
    meta: Meta = field(default_factory=Meta)
    visual: Visual = field(default_factory=Visual)


@dataclass
class Faction:
    id: str = field(default_factory=new_uuid)
    name: str = ""
    description: str = ""
    custom_tag_fields: List[CustomTagField] = field(default_factory=lambda: [CustomTagField() for _ in range(5)])
    meta: Meta = field(default_factory=Meta)
    visual: Visual = field(default_factory=Visual)


@dataclass
class Location:
    id: str = field(default_factory=new_uuid)
    name: str = ""
    description: str = ""
    custom_tag_fields: List[CustomTagField] = field(default_factory=lambda: [CustomTagField() for _ in range(5)])
    meta: Meta = field(default_factory=Meta)


@dataclass
class Story:
    id: str = field(default_factory=new_uuid)
    name: str = ""
    description: str = ""
    custom_tag_fields: List[CustomTagField] = field(default_factory=lambda: [CustomTagField() for _ in range(5)])
    meta: Meta = field(default_factory=Meta)


@dataclass
class Hook:
    id: str = field(default_factory=new_uuid)
    a_character_id: str = ""
    b_character_id: str = ""
    type: str = "positive"  # positive|conflict
    label: str = ""
    description: str = ""
    meta: Meta = field(default_factory=Meta)


@dataclass
class Goal:
    id: str = field(default_factory=new_uuid)
    character_id: str = ""
    status: str = "active"  # active|completed
    title: str = ""
    description: str = ""
    meta: Meta = field(default_factory=Meta)


@dataclass
class VersionEntry:
    id: str = field(default_factory=new_uuid)
    timestamp: str = field(default_factory=now_iso)
    kind: str = "manual"  # manual|auto
    author: str = "local"
    comment: str = ""
    state: dict = field(default_factory=dict)


def dataclass_to_dict(obj):
    return asdict(obj)

def dict_to_meta(d: dict) -> Meta:
    m = Meta()
    for k in ["created_at", "updated_at", "updated_by", "revision", "is_deleted"]:
        if k in d:
            setattr(m, k, d[k])
    return m

def dict_to_custom_fields(arr: list) -> List[CustomTagField]:
    out = []
    for item in (arr or []):
        out.append(CustomTagField(
            field_title=item.get("field_title", ""),
            values=list(item.get("values", []))
        ))
    while len(out) < 5:
        out.append(CustomTagField())
    return out[:5]

def dict_to_visual(d: dict) -> Visual:
    return Visual(x=float(d.get("x", 0.0)), y=float(d.get("y", 0.0)))


# =========================
# Project Container
# =========================

class Project:
    def __init__(self):
        self.schema_version = SCHEMA_VERSION
        self.characters: Dict[str, Character] = {}
        self.factions: Dict[str, Faction] = {}
        self.locations: Dict[str, Location] = {}
        self.stories: Dict[str, Story] = {}
        self.hooks: Dict[str, Hook] = {}
        self.goals: Dict[str, Goal] = {}

        self.character_custom_dicts: List[Set[str]] = [set() for _ in range(5)]
        self.faction_custom_dicts: List[Set[str]] = [set() for _ in range(5)]
        self.location_custom_dicts: List[Set[str]] = [set() for _ in range(5)]
        self.story_custom_dicts: List[Set[str]] = [set() for _ in range(5)]

        self.masters_dict: Set[str] = set()

        # Global titles for custom tag fields (synchronized inside one project file)
        # Separate namespaces for characters and factions:
        #   characters: field1..field5 titles shared across all characters
        #   factions:   field1..field5 titles shared across all factions
        self.character_custom_titles: List[str] = [""] * 5
        self.faction_custom_titles: List[str] = [""] * 5

        self.versions: List[VersionEntry] = []
        self.file_path: Optional[str] = None
        # Export settings
        self.export_font = "DejaVu Sans"
        self.export_defaults = {
            "name": True,
            "factions": True,
            "locations": True,
            "masters": False,
            "custom": [False] * 5,
            "public": False,
            "private": False,
            "goals": False,
            "hooks": False,
        }

    def build_name_index(self, exclude: Optional[Tuple[str, str]] = None) -> Dict[str, Tuple[str, str]]:
        idx = {}
        def add(obj_type: str, obj_id: str, name: str, is_deleted: bool):
            if is_deleted:
                return
            if exclude and exclude == (obj_type, obj_id):
                return
            nk = norm_key(name)
            if not nk:
                return
            idx[nk] = (obj_type, obj_id)

        for c in self.characters.values():
            add("character", c.id, c.name, c.meta.is_deleted)
        for f in self.factions.values():
            add("faction", f.id, f.name, f.meta.is_deleted)
        for l in self.locations.values():
            add("location", l.id, l.name, l.meta.is_deleted)
        for s in self.stories.values():
            add("story", s.id, s.name, s.meta.is_deleted)
        return idx

    def ensure_unique_name(self, obj_type: str, obj_id: str, new_name: str) -> Optional[str]:
        new_name = norm_spaces(new_name)
        if not new_name:
            return "Название обязательно."
        idx = self.build_name_index(exclude=(obj_type, obj_id))
        nk = norm_key(new_name)
        if nk in idx:
            t, _ = idx[nk]
            return f"Название уже занято ({t}). Переименуй."
        return None

    def rebuild_dictionaries(self):
        self.character_custom_dicts = [set() for _ in range(5)]
        self.faction_custom_dicts = [set() for _ in range(5)]
        self.location_custom_dicts = [set() for _ in range(5)]
        self.story_custom_dicts = [set() for _ in range(5)]
        self.masters_dict = set()

        for c in self.characters.values():
            if c.meta.is_deleted:
                continue
            for i in range(5):
                for v in c.custom_tag_fields[i].values:
                    v2 = norm_spaces(v)
                    if v2:
                        self.character_custom_dicts[i].add(v2)
            for m in c.masters:
                m2 = norm_spaces(m)
                if m2:
                    self.masters_dict.add(m2)

        for f in self.factions.values():
            if f.meta.is_deleted:
                continue
            for i in range(5):
                for v in f.custom_tag_fields[i].values:
                    v2 = norm_spaces(v)
                    if v2:
                        self.faction_custom_dicts[i].add(v2)

        for l in self.locations.values():
            if l.meta.is_deleted:
                continue
            for i in range(5):
                for v in l.custom_tag_fields[i].values:
                    v2 = norm_spaces(v)
                    if v2:
                        self.location_custom_dicts[i].add(v2)

        for s in self.stories.values():
            if s.meta.is_deleted:
                continue
            for i in range(5):
                for v in s.custom_tag_fields[i].values:
                    v2 = norm_spaces(v)
                    if v2:
                        self.story_custom_dicts[i].add(v2)

    # =========================
    # Custom field title sync
    # =========================

    def _ensure_custom_fields_len(self):
        # Defensive: make sure all objects have 5 custom fields
        for c in self.characters.values():
            if len(c.custom_tag_fields) < 5:
                c.custom_tag_fields += [CustomTagField() for _ in range(5 - len(c.custom_tag_fields))]
            elif len(c.custom_tag_fields) > 5:
                c.custom_tag_fields = c.custom_tag_fields[:5]
        for f in self.factions.values():
            if len(f.custom_tag_fields) < 5:
                f.custom_tag_fields += [CustomTagField() for _ in range(5 - len(f.custom_tag_fields))]
            elif len(f.custom_tag_fields) > 5:
                f.custom_tag_fields = f.custom_tag_fields[:5]

    def sync_custom_field_titles_from_objects(self):
        """Build global titles from existing objects and then enforce them everywhere."""
        self._ensure_custom_fields_len()

        # pick a title per index: first non-empty encountered (stable, deterministic)
        for i in range(5):
            title_c = ""
            for c in self.characters.values():
                if c.meta.is_deleted:
                    continue
                t = norm_spaces(getattr(c.custom_tag_fields[i], "field_title", ""))
                if t:
                    title_c = t
                    break
            if title_c or self.character_custom_titles[i]:
                self.character_custom_titles[i] = title_c or self.character_custom_titles[i]

            title_f = ""
            for f in self.factions.values():
                if f.meta.is_deleted:
                    continue
                t = norm_spaces(getattr(f.custom_tag_fields[i], "field_title", ""))
                if t:
                    title_f = t
                    break
            if title_f or self.faction_custom_titles[i]:
                self.faction_custom_titles[i] = title_f or self.faction_custom_titles[i]

        # enforce everywhere
        self.apply_custom_field_titles()

    def apply_custom_field_titles(self):
        """Apply current global titles to every character/faction."""
        self._ensure_custom_fields_len()
        for c in self.characters.values():
            if c.meta.is_deleted:
                continue
            for i in range(5):
                c.custom_tag_fields[i].field_title = self.character_custom_titles[i]
        for f in self.factions.values():
            if f.meta.is_deleted:
                continue
            for i in range(5):
                f.custom_tag_fields[i].field_title = self.faction_custom_titles[i]

    def set_character_custom_title(self, index: int, title: str):
        title = norm_spaces(title)
        if index < 0 or index >= 5:
            return
        self.character_custom_titles[index] = title
        # propagate to all characters (including currently edited)
        for c in self.characters.values():
            if c.meta.is_deleted:
                continue
            if len(c.custom_tag_fields) < 5:
                c.custom_tag_fields += [CustomTagField() for _ in range(5 - len(c.custom_tag_fields))]
            c.custom_tag_fields[index].field_title = title

    def set_faction_custom_title(self, index: int, title: str):
        title = norm_spaces(title)
        if index < 0 or index >= 5:
            return
        self.faction_custom_titles[index] = title
        for f in self.factions.values():
            if f.meta.is_deleted:
                continue
            if len(f.custom_tag_fields) < 5:
                f.custom_tag_fields += [CustomTagField() for _ in range(5 - len(f.custom_tag_fields))]
            f.custom_tag_fields[index].field_title = title


    def snapshot_state(self) -> dict:
        return {
            "schema_version": self.schema_version,
            "characters": [dataclass_to_dict(c) for c in self.characters.values()],
            "factions": [dataclass_to_dict(f) for f in self.factions.values()],
            "locations": [dataclass_to_dict(l) for l in self.locations.values()],
            "stories": [dataclass_to_dict(s) for s in self.stories.values()],
            "hooks": [dataclass_to_dict(h) for h in self.hooks.values()],
            "goals": [dataclass_to_dict(g) for g in self.goals.values()],
            "export_font": self.export_font,
            "export_defaults": self.export_defaults,
        }

    def load_state(self, state: dict):
        file_sv = int(state.get("schema_version", 1))
        self.schema_version = SCHEMA_VERSION

        self.characters = {}
        self.factions = {}
        self.locations = {}
        self.stories = {}
        self.hooks = {}
        self.goals = {}

        for d in state.get("characters", []):
            c = Character()
            c.id = d.get("id", new_uuid())
            c.name = d.get("name", "")
            c.story_public = d.get("story_public", "")
            c.story_private = d.get("story_private", "")
            c.factions = list(d.get("factions", []))
            c.locations = list(d.get("locations", []))
            c.stories = list(d.get("stories", [])) if file_sv >= 2 else list(d.get("stories", []))
            c.masters = list(d.get("masters", []))
            c.custom_tag_fields = dict_to_custom_fields(d.get("custom_tag_fields", []))
            c.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                c.meta.is_deleted = bool(d.get("is_deleted"))
            c.visual = dict_to_visual(d.get("visual", {}))
            self.characters[c.id] = c

        for d in state.get("factions", []):
            f = Faction()
            f.id = d.get("id", new_uuid())
            f.name = d.get("name", "")
            f.description = d.get("description", "")
            f.custom_tag_fields = dict_to_custom_fields(d.get("custom_tag_fields", []))
            f.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                f.meta.is_deleted = bool(d.get("is_deleted"))
            f.visual = dict_to_visual(d.get("visual", {}))
            self.factions[f.id] = f

        for d in state.get("locations", []):
            l = Location()
            l.id = d.get("id", new_uuid())
            l.name = d.get("name", "")
            l.description = d.get("description", "")
            l.custom_tag_fields = dict_to_custom_fields(d.get("custom_tag_fields", []))
            l.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                l.meta.is_deleted = bool(d.get("is_deleted"))
            self.locations[l.id] = l

        for d in state.get("stories", []):
            s = Story()
            s.id = d.get("id", new_uuid())
            s.name = d.get("name", "")
            s.description = d.get("description", "")
            s.custom_tag_fields = dict_to_custom_fields(d.get("custom_tag_fields", []))
            s.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                s.meta.is_deleted = bool(d.get("is_deleted"))
            self.stories[s.id] = s

        for d in state.get("hooks", []):
            h = Hook()
            h.id = d.get("id", new_uuid())
            h.a_character_id = d.get("a_character_id", "")
            h.b_character_id = d.get("b_character_id", "")
            h.type = d.get("type", "positive")
            h.label = d.get("label", "")
            h.description = d.get("description", "")
            h.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                h.meta.is_deleted = bool(d.get("is_deleted"))
            self.hooks[h.id] = h

        for d in state.get("goals", []):
            g = Goal()
            g.id = d.get("id", new_uuid())
            g.character_id = d.get("character_id", "")
            g.status = d.get("status", "active")
            if g.status not in ("active", "completed"):
                g.status = "active"
            g.title = d.get("title", "")
            g.description = d.get("description", "")
            g.meta = dict_to_meta(d.get("meta", {}))
            if "is_deleted" in d:
                g.meta.is_deleted = bool(d.get("is_deleted"))
            self.goals[g.id] = g

        # export settings
        ef = state.get("export_font")
        if isinstance(ef, str) and ef.strip():
            self.export_font = ef.strip()
        ed = state.get("export_defaults")
        if isinstance(ed, dict):
            base = {
                "name": True,
                "factions": True,
                "locations": True,
                "masters": False,
                "custom": [False] * 5,
                "public": False,
                "private": False,
                "goals": False,
                "hooks": False,
            }
            base.update(ed)
            if not isinstance(base.get("custom"), list) or len(base.get("custom", [])) < 5:
                base["custom"] = (base.get("custom") or []) + [False] * 5
            base["custom"] = list(base["custom"])[:5]
            self.export_defaults = base

        # Synchronize custom field titles inside this project file
        self.sync_custom_field_titles_from_objects()

        self.rebuild_dictionaries()

    def to_file_dict(self) -> dict:
        return {
            "schema_version": SCHEMA_VERSION,
            "current": self.snapshot_state(),
            "versions": [dataclass_to_dict(v) for v in self.versions],
        }

    def from_file_dict(self, d: dict):
        file_sv = int(d.get("schema_version", 1))
        current = d.get("current", {})
        if "schema_version" not in current:
            current["schema_version"] = file_sv
        self.load_state(current)

        self.versions = []
        for vd in d.get("versions", []):
            ve = VersionEntry()
            ve.id = vd.get("id", new_uuid())
            ve.timestamp = vd.get("timestamp", now_iso())
            ve.kind = vd.get("kind", "manual")
            ve.author = vd.get("author", "local")
            ve.comment = vd.get("comment", "")
            ve.state = vd.get("state", {})
            self.versions.append(ve)

    def save_to_path(self, path: str, kind: str = "manual", author: str = "local", comment: str = ""):
        ve = VersionEntry(kind=kind, author=author, comment=comment, state=self.snapshot_state())
        self.versions.append(ve)

        payload = self.to_file_dict()
        tmp_path = path + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        shutil.move(tmp_path, path)
        self.file_path = path

    def open_from_path(self, path: str):
        with open(path, "r", encoding="utf-8") as f:
            d = json.load(f)
        self.from_file_dict(d)
        self.file_path = path

    def get_object_by_name(self, name: str) -> Optional[Tuple[str, str]]:
        nk = norm_key(name)
        if not nk:
            return None
        for c in self.characters.values():
            if not c.meta.is_deleted and norm_key(c.name) == nk:
                return ("character", c.id)
        for f in self.factions.values():
            if not f.meta.is_deleted and norm_key(f.name) == nk:
                return ("faction", f.id)
        for l in self.locations.values():
            if not l.meta.is_deleted and norm_key(l.name) == nk:
                return ("location", l.id)
        for s in self.stories.values():
            if not s.meta.is_deleted and norm_key(s.name) == nk:
                return ("story", s.id)
        return None

    def alive_characters(self) -> List[Character]:
        return [c for c in self.characters.values() if not c.meta.is_deleted]

    def alive_factions(self) -> List['Faction']:
        return [f for f in self.factions.values() if not f.meta.is_deleted]

    def alive_locations(self) -> List['Location']:
        return [l for l in self.locations.values() if not l.meta.is_deleted]

    def alive_stories(self) -> List['Story']:
        return [s for s in self.stories.values() if not s.meta.is_deleted]

    def goals_for_character(self, char_id: str) -> List[Goal]:
        gs = [g for g in self.goals.values() if (not g.meta.is_deleted and g.character_id == char_id)]
        gs.sort(key=lambda x: (0 if x.status == "active" else 1, x.title.casefold(), x.id))
        return gs

    def hooks_for_character(self, char_id: str) -> List[Hook]:
        hs = [
            h for h in self.hooks.values()
            if (not h.meta.is_deleted and (h.a_character_id == char_id or h.b_character_id == char_id))
        ]
        hs.sort(key=lambda x: (x.type, x.label.casefold(), x.id))
        return hs



# =========================
# Online Sync (server + launcher)
# =========================

@dataclass
class OnlineContext:
    mode: str = "offline"  # offline|online
    server_url: str = ""
    game_id: str = ""
    access_token: str = ""
    activation_key: str = ""
    username: str = "local"

class SyncError(Exception):
    pass

def _http_headers(activation_key: str, access_token: str) -> Dict[str, str]:
    h = {"Content-Type": "application/json"}
    if activation_key:
        h["X-Activation-Key"] = activation_key
        h["X-SW-Enc"] = "1"
    if access_token:
        h["Authorization"] = f"Bearer {access_token}"
    return h

def server_url_to_ws(server_url: str) -> str:
    s = (server_url or "").strip()
    if s.startswith("https://"):
        return "wss://" + s[len("https://"):]
    if s.startswith("http://"):
        return "ws://" + s[len("http://"):]
    # assume already ws/wss
    return s

def project_to_server_state(project: 'Project') -> Dict[str, Any]:
    # Store entities by id to allow small ops and conflict-free merges.
    # IMPORTANT: server-side ops can "soft_delete" entities by setting top-level `is_deleted`.
    # Our editor keeps deletion in `meta.is_deleted`, so we mirror it to `is_deleted` for compatibility.
    def _pack(ent_obj) -> Dict[str, Any]:
        d = dataclass_to_dict(ent_obj)
        try:
            md = d.get("meta")
            if isinstance(md, dict) and "is_deleted" in md:
                d["is_deleted"] = bool(md.get("is_deleted"))
        except Exception:
            pass
        return d

    return {
        "schema_version": SCHEMA_VERSION,
        "characters": {cid: _pack(c) for cid, c in project.characters.items()},
        "factions": {fid: _pack(f) for fid, f in project.factions.items()},
        "locations": {lid: _pack(l) for lid, l in project.locations.items()},
        "stories": {sid: _pack(s) for sid, s in project.stories.items()},
        "hooks": {hid: _pack(h) for hid, h in project.hooks.items()},
        "goals": {gid: _pack(g) for gid, g in project.goals.items()},
        # keep custom field title sync stable across clients
        "custom_titles": {
            "characters": list(project.character_custom_titles),
            "factions": list(project.faction_custom_titles),
        },
        "export_settings": {
            "font": project.export_font,
            "defaults": project.export_defaults,
        },
    }

def server_state_to_editor_current(state: Dict[str, Any]) -> Dict[str, Any]:
    # Convert server dict-of-dicts into editor's list format for Project.load_state()
    def _vals(key: str) -> List[dict]:
        d = state.get(key) or {}
        if isinstance(d, dict):
            return list(d.values())
        if isinstance(d, list):
            return d
        return []

    cur = {
        "schema_version": int(state.get("schema_version", SCHEMA_VERSION)),
        "characters": _vals("characters"),
        "factions": _vals("factions"),
        "locations": _vals("locations"),
        "stories": _vals("stories"),
        "hooks": _vals("hooks"),
        "goals": _vals("goals"),
    }
    # export settings (optional)
    ex = state.get("export_settings") or {}
    if isinstance(ex, dict):
        if isinstance(ex.get("font"), str):
            cur["export_font"] = ex.get("font")
        if isinstance(ex.get("defaults"), dict):
            cur["export_defaults"] = ex.get("defaults")
    return cur

class SyncApiClient:
    def __init__(self, ctx: OnlineContext, timeout: int = 20):
        self.ctx = ctx
        self.timeout = timeout

    def _base(self) -> str:
        return (self.ctx.server_url or "").rstrip("/")

    def _url(self, path: str) -> str:
        if not path.startswith("/"):
            path = "/" + path
        return self._base() + path

    def get_snapshot(self) -> Tuple[Dict[str, Any], int]:
        if not self.ctx.game_id:
            raise SyncError("game_id missing")
        # Server endpoint: GET /sync/snapshot/{game_id}
        url = self._url(f"/sync/snapshot/{self.ctx.game_id}")
        r = requests.get(url, headers=_http_headers(self.ctx.activation_key, self.ctx.access_token), timeout=self.timeout)
        if r.status_code >= 400:
            msg = r.text
            try:
                data = r.json()
                if self.ctx.activation_key:
                    data = _http_decrypt_payload(self.ctx.activation_key, data)
                if isinstance(data, dict) and data.get("detail"):
                    msg = data.get("detail")
                else:
                    msg = str(data)
            except Exception:
                pass
            raise SyncError(f"{r.status_code}: {msg}")
        d = r.json()
        if self.ctx.activation_key:
            try:
                d = _http_decrypt_payload(self.ctx.activation_key, d)
            except Exception:
                pass
        return (d.get("state") or {}, int(d.get("revision", 0)))

    def apply_ops(self, base_revision: int, ops: List[Dict[str, Any]]) -> int:
        url = self._url("/sync/apply_ops")
        payload = {"game_id": self.ctx.game_id, "base_revision": int(base_revision), "ops": ops}
        if self.ctx.activation_key:
            try:
                payload = _http_encrypt_payload(self.ctx.activation_key, payload)
            except Exception:
                pass
        r = requests.post(url, json=payload, headers=_http_headers(self.ctx.activation_key, self.ctx.access_token), timeout=self.timeout)
        if r.status_code == 409:
            raise SyncError("conflict")
        if r.status_code >= 400:
            msg = r.text
            try:
                data = r.json()
                if self.ctx.activation_key:
                    data = _http_decrypt_payload(self.ctx.activation_key, data)
                if isinstance(data, dict) and data.get("detail"):
                    msg = data.get("detail")
                else:
                    msg = str(data)
            except Exception:
                pass
            raise SyncError(f"{r.status_code}: {msg}")
        d = r.json()
        if self.ctx.activation_key:
            try:
                d = _http_decrypt_payload(self.ctx.activation_key, d)
            except Exception:
                pass
        return int(d.get("new_revision", base_revision))

    def save_version(self, comment: str = ""):
        # Optional: server-side version snapshot
        url = self._url("/versions/save")
        payload = {"game_id": self.ctx.game_id, "comment": comment or ""}
        if self.ctx.activation_key:
            try:
                payload = _http_encrypt_payload(self.ctx.activation_key, payload)
            except Exception:
                pass
        r = requests.post(url, json=payload, headers=_http_headers(self.ctx.activation_key, self.ctx.access_token), timeout=self.timeout)
        if r.status_code >= 400:
            msg = r.text
            try:
                data = r.json()
                if self.ctx.activation_key:
                    data = _http_decrypt_payload(self.ctx.activation_key, data)
                if isinstance(data, dict) and data.get("detail"):
                    msg = data.get("detail")
                else:
                    msg = str(data)
            except Exception:
                pass
            raise SyncError(f"{r.status_code}: {msg}")
        d = r.json()
        if self.ctx.activation_key:
            try:
                d = _http_decrypt_payload(self.ctx.activation_key, d)
            except Exception:
                pass
        return d

def _stable_json(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, sort_keys=True, separators=(",", ":"))

def diff_state_to_ops(old: Dict[str, Any], new: Dict[str, Any]) -> List[Dict[str, Any]]:
    ops: List[Dict[str, Any]] = []
    # schema version
    if int(old.get("schema_version", SCHEMA_VERSION)) != int(new.get("schema_version", SCHEMA_VERSION)):
        ops.append({"op": "set", "path": "schema_version", "value": int(new.get("schema_version", SCHEMA_VERSION))})

    # custom titles
    if _stable_json(old.get("custom_titles")) != _stable_json(new.get("custom_titles")):
        ops.append({"op": "set", "path": "custom_titles", "value": new.get("custom_titles") or {}})

    for col in ["characters", "factions", "locations", "stories", "hooks", "goals"]:
        ocol = old.get(col) or {}
        ncol = new.get(col) or {}
        if not isinstance(ocol, dict):
            ocol = {}
        if not isinstance(ncol, dict):
            ncol = {}

        # deletions (hard delete from server state) are rare in editor, but keep it safe:
        for eid in set(ocol.keys()) - set(ncol.keys()):
            ops.append({"op": "delete", "path": f"{col}.{eid}"})

        # upserts/changes
        for eid, ent in ncol.items():
            if eid not in ocol or _stable_json(ocol.get(eid)) != _stable_json(ent):
                ops.append({"op": "upsert_entity", "collection": col, "id": eid, "value": ent})
    return ops

class RealtimeSync(QtCore.QObject):
    statusChanged = QtCore.Signal(str)
    stateReceived = QtCore.Signal(dict, int)   # state, revision
    remoteApplied = QtCore.Signal()
    errorOccurred = QtCore.Signal(str)

    def __init__(self, api: SyncApiClient, ctx: OnlineContext, parent=None):
        super().__init__(parent)
        self.api = api
        self.ctx = ctx

        self.ws = QWebSocket()
        self.ws.connected.connect(self._on_connected)
        self.ws.disconnected.connect(self._on_disconnected)
        self.ws.textMessageReceived.connect(self._on_text)

        self._connected = False
        self._subscribed = False
        self._my_role = None
        self._last_server_sv = None

        # Authoritative server baseline at `revision`
        self.revision: int = 0
        self.server_state: Dict[str, Any] = {}

        # Baseline + optimistic local changes that were sent but not yet confirmed.
        self.shadow_state: Dict[str, Any] = {}

        # In-flight batch (one at a time). This keeps multi-client edits from "eating" each other.
        self._inflight_ops: Optional[List[Dict[str, Any]]] = None
        self._inflight_hash: Optional[str] = None
        self._inflight_base_rev: Optional[int] = None

        # Debounced local change flush
        self._dirty = False
        self._debounce = QtCore.QTimer(self)
        self._debounce.setInterval(250)
        self._debounce.setSingleShot(True)
        self._debounce.timeout.connect(self._flush_if_needed)

    def start(self):
        # Build ws URL with query params (server supports token + activation_key in query).
        base_ws = server_url_to_ws(self.ctx.server_url).rstrip("/")
        qs = {}
        if self.ctx.access_token:
            qs["token"] = self.ctx.access_token
        if self.ctx.activation_key:
            qs["activation_key"] = self.ctx.activation_key
        q = urllib.parse.urlencode(qs)
        url = f"{base_ws}/ws"
        if q:
            url += "?" + q

        self.statusChanged.emit("Подключение к серверу…")
        self.ws.open(QNetworkRequest(QtCore.QUrl(url)))
        # Polling starts only after initial subscribe state is applied (prevents overwriting server state on startup)

    def stop(self):
        try:
            self._poll.stop()
        except Exception:
            pass
        try:
            self.ws.close()
        except Exception:
            pass

    def mark_dirty(self):
        # Don't spam the network. Debounce.
        self._dirty = True
        self._debounce.start()

    def flush_full(self, project: 'Project'):
        """
        Force a flush (used on app close / manual save).
        IMPORTANT: do NOT "set whole collections" (that stomps others). Only send a diff.
        """
        if not self._subscribed:
            return
        try:
            new_state = project_to_server_state(project)
            base = self.shadow_state if isinstance(self.shadow_state, dict) else (self.server_state or {})
            ops = diff_state_to_ops(base, new_state)
            if not ops:
                return
            # REST path is fine here: it still broadcasts to other clients.
            self.revision = self.api.apply_ops(self.revision, ops)
            # committed
            self.server_state = apply_ops_local(self.server_state, ops)
            self.shadow_state = self._clone_state(self.server_state)
            self._clear_inflight()
        except SyncError as e:
            # If conflict during close-save, just rebase once and try again.
            if str(e) == "conflict":
                try:
                    snap, rev = self.api.get_snapshot()
                    self.server_state = snap or {}
                    self.shadow_state = self._clone_state(self.server_state)
                    self.revision = int(rev)
                    new_state = project_to_server_state(project)
                    ops = diff_state_to_ops(self.shadow_state, new_state)
                    if ops:
                        self.revision = self.api.apply_ops(self.revision, ops)
                        self.server_state = apply_ops_local(self.server_state, ops)
                        self.shadow_state = self._clone_state(self.server_state)
                        self._clear_inflight()
                except Exception as e2:
                    self.errorOccurred.emit(str(e2))
            else:
                self.errorOccurred.emit(str(e))
        except Exception as e:
            self.errorOccurred.emit(str(e))

    # ---- internal ----
    def _clone_state(self, st: Dict[str, Any]) -> Dict[str, Any]:
        # Deep clone to avoid shared nested dict references (ops mutate in-place).
        try:
            return json.loads(json.dumps(st or {}, ensure_ascii=False))
        except Exception:
            return {}

    def _hash_ops(self, ops: List[Dict[str, Any]]) -> str:
        return _stable_json(ops)

    def _clear_inflight(self):
        self._inflight_ops = None
        self._inflight_hash = None
        self._inflight_base_rev = None

    def _on_connected(self):
        self._connected = True
        self.statusChanged.emit("Подключено. Подписка на игру…")
        msg = {"type": "subscribe", "game_id": self.ctx.game_id}
        self.ws.sendTextMessage(json.dumps(msg, ensure_ascii=False))

    def _on_disconnected(self):
        self._connected = False
        self._subscribed = False
        self._my_role = None
        self._last_server_sv = None
        self._clear_inflight()
        self.statusChanged.emit("Отключено от сервера.")

    def _on_text(self, txt: str):
        try:
            msg = json.loads(txt or "{}")
        except Exception:
            return

        typ = msg.get("type")

        if typ == "subscribed":
            self._subscribed = True
            st = msg.get("state") or {}
            self.server_state = st
            self.shadow_state = self._clone_state(st)  # shallow ok, we only replace whole entities
            self.revision = int(msg.get("revision", 0))
            self._clear_inflight()
            self.statusChanged.emit("Синхронизация активна.")
            self.stateReceived.emit(self.shadow_state, self.revision)
            return

        if typ == "ack":
            # Server accepted our last batch. Commit it immediately (no "random client wins").
            new_rev = int(msg.get("new_revision", self.revision))
            if self._inflight_ops:
                # Commit optimistic changes into authoritative baseline.
                self.server_state = apply_ops_local(self.server_state, self._inflight_ops)
                self.revision = new_rev
                self.shadow_state = self._clone_state(self.server_state)
                self._clear_inflight()
                # If user edited more while we were waiting, flush again.
                if self._dirty:
                    self._debounce.start()
            else:
                self.revision = max(self.revision, new_rev)
            return

        if typ == "conflict":
            # Somebody else wrote first. Rebase local state onto fresh snapshot and re-send.
            self.statusChanged.emit("Конфликт версий. Перебазирую изменения…")
            try:
                local_state = {}
                if hasattr(self, "get_current_state"):
                    try:
                        local_state = self.get_current_state() or {}
                    except Exception:
                        local_state = {}
                snap, rev = self.api.get_snapshot()
                snap = snap or {}
                rev = int(rev)

                # Throw away the in-flight batch: server rejected it.
                self._clear_inflight()

                # Merge snapshot + current local state, so UI doesn't "eat" edits.
                reapply_ops = diff_state_to_ops(snap, local_state) if local_state else []
                merged = apply_ops_local(self._clone_state(snap), reapply_ops) if reapply_ops else self._clone_state(snap)

                # Baseline is snapshot; shadow is what UI shows.
                self.server_state = self._clone_state(snap)
                self.shadow_state = self._clone_state(merged)
                self.revision = rev

                self.stateReceived.emit(self.shadow_state, self.revision)

                # If we had anything local to reapply, send it now on top of snapshot.
                if reapply_ops:
                    self._send_ops_now(reapply_ops)
            except Exception as e:
                self.errorOccurred.emit(str(e))
            return

        if typ == "event":
            ops = msg.get("ops") or []
            base_rev = int(msg.get("base_revision", -1))
            new_rev = int(msg.get("new_revision", self.revision))
            ops_hash = self._hash_ops(ops)

            # If this is an echo of our inflight batch, do not double-apply into shadow.
            if self._inflight_hash and ops_hash == self._inflight_hash and base_rev == int(self._inflight_base_rev or -2):
                # Commit to authoritative baseline (shadow already has it).
                self.server_state = apply_ops_local(self.server_state, ops)
                self.revision = new_rev
                self.shadow_state = self._clone_state(self.server_state)
                self._clear_inflight()
                return

            try:
                self.server_state = apply_ops_local(self.server_state, ops)
                self.shadow_state = apply_ops_local(self.shadow_state, ops)
                self.revision = new_rev
                self.remoteApplied.emit()
            except Exception as e:
                self.errorOccurred.emit(str(e))
            return

        if typ == "error":
            self.errorOccurred.emit(str(msg.get("detail", "error")))
            return

    def _send_ops_now(self, ops: List[Dict[str, Any]]):
        if not self._subscribed:
            return
        if self._inflight_ops:
            # Can't send right now; will be sent by the debounced flush.
            self._dirty = True
            self._debounce.start()
            return
        if not ops:
            return

        payload = {"type": "event", "base_revision": int(self.revision), "ops": ops}
        try:
            self._inflight_ops = ops
            self._inflight_hash = self._hash_ops(ops)
            self._inflight_base_rev = int(self.revision)
            # Optimistically apply into shadow (UI already has it, but baseline for next diff must include it)
            self.shadow_state = apply_ops_local(self.shadow_state, ops)
            self.ws.sendTextMessage(json.dumps(payload, ensure_ascii=False))
        except Exception as e:
            self._clear_inflight()
            self.errorOccurred.emit(str(e))


def _poll_local_state(self):
    if not self._subscribed or self.get_current_state is None or self._applying_remote:
        return
    try:
        st = self.get_current_state() or {}
        h = hashlib.sha256(_stable_json(st).encode("utf-8")).hexdigest()
        if h != self._last_local_hash:
            self._last_local_hash = h
            self.mark_dirty()
            log_event("CRDT_POLL_DIRTY", hash=h[:12])
    except Exception as e:
        log_event("CRDT_POLL_ERR", err=str(e))

    def _flush_if_needed(self):
        if not self._dirty:
            return
        if not self._connected or not self._subscribed:
            return
        if not self._subscribed:
            return
        if self._inflight_ops:
            # Wait for ack/echo and try again.
            self._debounce.start()
            return
        if not hasattr(self, "get_current_state"):
            return

        try:
            new_state = self.get_current_state() or {}
        except Exception as e:
            self.errorOccurred.emit(str(e))
            return

        base = self.shadow_state if isinstance(self.shadow_state, dict) else (self.server_state or {})
        ops = diff_state_to_ops(base, new_state)
        if not ops:
            self._dirty = False
            return

        self._dirty = False
        self._send_ops_now(ops)

class LocalOpError(Exception):
    pass

def _set_path_local(obj: Any, path: str, value: Any):
    parts = [p for p in (path or "").split(".") if p]
    if not parts:
        raise LocalOpError("empty path")
    cur = obj
    for p in parts[:-1]:
        if not isinstance(cur, dict):
            raise LocalOpError(f"path invalid at '{p}'")
        if p not in cur or not isinstance(cur[p], dict):
            cur[p] = {}
        cur = cur[p]
    if not isinstance(cur, dict):
        raise LocalOpError("path invalid (leaf parent not dict)")
    cur[parts[-1]] = value

def _del_path_local(obj: Any, path: str):
    parts = [p for p in (path or "").split(".") if p]
    if not parts:
        return
    cur = obj
    for p in parts[:-1]:
        if not isinstance(cur, dict):
            return
        cur = cur.get(p)
        if cur is None:
            return
    if isinstance(cur, dict):
        cur.pop(parts[-1], None)

def apply_ops_local(state: Dict[str, Any], ops: List[Dict[str, Any]]) -> Dict[str, Any]:
    # Same semantics as server ops.py (kept here to avoid a dependency on server files)
    st = state if isinstance(state, dict) else {}
    for o in ops or []:
        typ = o.get("op")
        if typ == "set":
            _set_path_local(st, o.get("path", ""), o.get("value"))
        elif typ == "delete":
            _del_path_local(st, o.get("path", ""))
        elif typ == "upsert_entity":
            col = o.get("collection")
            eid = o.get("id")
            val = o.get("value") or {}
            if col not in st or not isinstance(st.get(col), dict):
                st[col] = {}
            st[col][eid] = val
        elif typ == "soft_delete_entity":
            col = o.get("collection")
            eid = o.get("id")
            if col in st and isinstance(st.get(col), dict) and eid in st[col]:
                ent = st[col][eid]
                if isinstance(ent, dict):
                    ent["is_deleted"] = True
                    md = ent.get("meta")
                    if isinstance(md, dict):
                        md["is_deleted"] = True
        else:
            raise LocalOpError(f"Unknown op: {typ}")
    return st


# =========================
# Hyperlink TextEdit
# =========================

class LinkHighlighter(QtGui.QSyntaxHighlighter):
    def __init__(self, doc: QtGui.QTextDocument, resolver):
        super().__init__(doc)
        self.resolver = resolver

        self.ok_format = QtGui.QTextCharFormat()
        self.ok_format.setForeground(QtGui.QColor("#2a62ff"))
        self.ok_format.setFontUnderline(True)

        self.bad_format = QtGui.QTextCharFormat()
        self.bad_format.setForeground(QtGui.QColor("#cc3333"))
        self.bad_format.setFontUnderline(True)

    def highlightBlock(self, text: str):
        for m in LINK_RE.finditer(text):
            inner = m.group(1)
            if words_count(inner) > 4:
                continue
            obj = self.resolver(inner)
            fmt = self.ok_format if obj is not None else self.bad_format
            start, end = m.span()
            self.setFormat(start, end - start, fmt)


class HyperTextEdit(QtWidgets.QTextEdit):
    linkActivated = QtCore.Signal(str)

    def __init__(self, resolver, parent=None):
        super().__init__(parent)
        self.resolver = resolver
        self.highlighter = LinkHighlighter(self.document(), self.resolver)
        self.setAcceptRichText(False)

    def mouseDoubleClickEvent(self, event: QtGui.QMouseEvent):
        cursor = self.cursorForPosition(event.position().toPoint())
        block = cursor.block()
        block_text = block.text()
        pos_in_block = cursor.position() - block.position()
        hit = extract_link_at_pos(block_text, pos_in_block)
        if hit:
            _, _, inner = hit
            if words_count(inner) <= 4 and self.resolver(inner) is not None:
                self.linkActivated.emit(inner)
                return
        super().mouseDoubleClickEvent(event)


# =========================
# MultiSelect Widgets
# =========================

class CheckListWidget(QtWidgets.QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSelectionMode(QtWidgets.QAbstractItemView.NoSelection)

    def set_items(self, items: List[str], checked: Set[str]):
        self.clear()
        for it in sorted(items, key=lambda s: s.casefold()):
            item = QtWidgets.QListWidgetItem(it)
            item.setFlags(item.flags() | QtCore.Qt.ItemIsUserCheckable)
            item.setCheckState(QtCore.Qt.Checked if it in checked else QtCore.Qt.Unchecked)
            self.addItem(item)

    def checked_items(self) -> List[str]:
        out = []
        for i in range(self.count()):
            item = self.item(i)
            if item.checkState() == QtCore.Qt.Checked:
                out.append(item.text())
        return out


# =========================
# Field / Graphics Scene
# =========================

class CharacterNodeItem(QtWidgets.QGraphicsEllipseItem):
    def __init__(self, char_id: str, label: str):
        super().__init__(-18, -18, 36, 36)
        self.char_id = char_id
        self.setBrush(QtGui.QBrush(QtGui.QColor("#2a3140")))
        self.setPen(QtGui.QPen(QtGui.QColor("#4b556b"), 1.5))

        self.setFlag(QtWidgets.QGraphicsItem.ItemIsMovable, True)
        self.setFlag(QtWidgets.QGraphicsItem.ItemIsSelectable, False)
        self.setFlag(QtWidgets.QGraphicsItem.ItemSendsGeometryChanges, True)
        self.setCacheMode(QtWidgets.QGraphicsItem.DeviceCoordinateCache)
        self.setZValue(10)

        self.text = QtWidgets.QGraphicsSimpleTextItem(label, self)
        self.text.setPos(-18, 22)
        self.text.setBrush(QtGui.QBrush(QtGui.QColor("#e6e6e6")))

    def itemChange(self, change, value):
        # Notify after the position is actually changed (avoids feedback/jitter while dragging)
        if change == QtWidgets.QGraphicsItem.ItemPositionHasChanged:
            scene = self.scene()
            if hasattr(scene, "on_node_moved"):
                scene.on_node_moved()
        return super().itemChange(change, value)


class FactionContainerItem(QtWidgets.QGraphicsEllipseItem):
    def __init__(self, faction_id: str, label: str):
        super().__init__(-90, -90, 180, 180)
        self.faction_id = faction_id
        self.setBrush(QtGui.QBrush(QtGui.QColor(60, 120, 255, 35)))
        self.setPen(QtGui.QPen(QtGui.QColor("#2b64ff"), 2.0))

        self.setFlag(QtWidgets.QGraphicsItem.ItemIsMovable, True)
        self.setFlag(QtWidgets.QGraphicsItem.ItemSendsGeometryChanges, True)
        self.setCacheMode(QtWidgets.QGraphicsItem.DeviceCoordinateCache)
        self.setZValue(1)

        self.text = QtWidgets.QGraphicsSimpleTextItem(label, self)
        self.text.setPos(-85, -110)
        self.text.setBrush(QtGui.QBrush(QtGui.QColor("#cfd6e6")))

    def itemChange(self, change, value):
        # Notify after the position is actually changed (avoids feedback/jitter while dragging)
        if change == QtWidgets.QGraphicsItem.ItemPositionHasChanged:
            scene = self.scene()
            if hasattr(scene, "on_node_moved"):
                scene.on_node_moved()
        return super().itemChange(change, value)


class HookLineItem(QtWidgets.QGraphicsPathItem):
    def __init__(self, hook_id: str, label: str, kind: str):
        super().__init__()
        self.hook_id = hook_id
        self.kind = kind
        pen = QtGui.QPen(QtGui.QColor("#2f7d32") if kind == "positive" else QtGui.QColor("#b23a3a"), 2.0)
        if kind == "conflict":
            pen.setStyle(QtCore.Qt.DashLine)
        self.setPen(pen)
        self.setZValue(5)
        self.setAcceptHoverEvents(True)

        self.text = QtWidgets.QGraphicsSimpleTextItem(label, None)
        self.text.setZValue(6)
        self.text.setBrush(QtGui.QBrush(QtGui.QColor("#e6e6e6")))

    def set_label_pos(self, x: float, y: float):
        self.text.setPos(x, y)

    def hoverEnterEvent(self, event):
        scene = self.scene()
        if hasattr(scene, "set_pair_hover") and hasattr(self, "pair_key"):
            scene.set_pair_hover(self.pair_key, True)
        super().hoverEnterEvent(event)

    def hoverLeaveEvent(self, event):
        scene = self.scene()
        if hasattr(scene, "set_pair_hover") and hasattr(self, "pair_key"):
            scene.set_pair_hover(self.pair_key, False)
        super().hoverLeaveEvent(event)

    def mousePressEvent(self, event: QtWidgets.QGraphicsSceneMouseEvent):
        if event.button() == QtCore.Qt.LeftButton:
            scene = self.scene()
            if hasattr(scene, "hook_clicked"):
                scene.hook_clicked(self.hook_id)
                event.accept()
                return
        super().mousePressEvent(event)


class FieldScene(QtWidgets.QGraphicsScene):
    def __init__(self, project: 'Project', parent=None):
        super().__init__(parent)
        self.project = project
        self.char_items: Dict[str, CharacterNodeItem] = {}
        self.faction_items: Dict[str, FactionContainerItem] = {}
        self.hook_items: Dict[str, HookLineItem] = {}
        self.aff_lines: List[QtWidgets.QGraphicsLineItem] = []
        self.pair_hover: Set[Tuple[str, str]] = set()
        self._suppress = False

        self.selection_enabled = False

        self.on_open_character = None
        self.on_open_hooks = None
        self.on_open_goals = None
        self.on_open_hook_by_id = None

        # NEW: create callbacks for background context menu
        self.on_create_character_at = None   # (QPointF)
        self.on_create_faction_at = None     # (QPointF)


        # NEW: faction/character actions from node UI
        self.on_open_faction = None          # (faction_id)
        self.on_delete_character = None      # (char_id)
        self.on_delete_faction = None        # (faction_id)
    def set_selection_enabled(self, enabled: bool):
        enabled = bool(enabled)
        if self.selection_enabled == enabled:
            return
        self.selection_enabled = enabled
        # toggle selectable flags
        for it in self.char_items.values():
            try:
                it.setFlag(QtWidgets.QGraphicsItem.ItemIsSelectable, enabled)
            except Exception:
                pass
        for it in self.faction_items.values():
            try:
                it.setFlag(QtWidgets.QGraphicsItem.ItemIsSelectable, enabled)
            except Exception:
                pass
        if not enabled:
            try:
                self.clearSelection()
            except Exception:
                pass

    def sync_visuals_to_project(self):
        for cid, it in self.char_items.items():
            c = self.project.characters.get(cid)
            if c and not c.meta.is_deleted:
                p = it.pos()
                c.visual.x, c.visual.y = float(p.x()), float(p.y())

        for fid, it in self.faction_items.items():
            f = self.project.factions.get(fid)
            if f and not f.meta.is_deleted:
                p = it.pos()
                f.visual.x, f.visual.y = float(p.x()), float(p.y())

    def rebuild(self):
        self.clear()
        self.char_items.clear()
        self.faction_items.clear()
        self.hook_items.clear()
        self.aff_lines.clear()
        self.pair_hover.clear()

        for f in self.project.alive_factions():
            it = FactionContainerItem(f.id, f.name)
            it.setPos(f.visual.x, f.visual.y)
            self.addItem(it)
            self.faction_items[f.id] = it

        for c in self.project.alive_characters():
            it = CharacterNodeItem(c.id, c.name)
            it.setPos(c.visual.x, c.visual.y)
            self.addItem(it)
            self.char_items[c.id] = it
            try:
                it.setFlag(QtWidgets.QGraphicsItem.ItemIsSelectable, self.selection_enabled)
            except Exception:
                pass

        # factions selectable only in selection mode
        for it in self.faction_items.values():
            try:
                it.setFlag(QtWidgets.QGraphicsItem.ItemIsSelectable, self.selection_enabled)
            except Exception:
                pass

        self.draw_all_lines()

    def set_pair_hover(self, pair_key: Tuple[str, str], hovered: bool):
        if hovered:
            self.pair_hover.add(pair_key)
        else:
            self.pair_hover.discard(pair_key)
        self.draw_all_lines()

    def on_node_moved(self):
        if self._suppress:
            return
        self.sync_visuals_to_project()
        self.draw_all_lines()
        # notify main window (for online sync / dirty tracking)
        cb = getattr(self, "on_project_changed", None)
        if cb:
            try:
                cb()
            except Exception:
                pass

    def draw_all_lines(self):
        for l in self.aff_lines:
            self.removeItem(l)
        self.aff_lines = []
        for hid, item in list(self.hook_items.items()):
            self.removeItem(item)
            self.removeItem(item.text)
        self.hook_items = {}

        # affiliations (character -> factions)
        for c in self.project.alive_characters():
            c_it = self.char_items.get(c.id)
            if not c_it:
                continue
            for fid in c.factions:
                f_it = self.faction_items.get(fid)
                if not f_it:
                    continue
                line = QtWidgets.QGraphicsLineItem(QtCore.QLineF(c_it.pos(), f_it.pos()))
                line.setPen(QtGui.QPen(QtGui.QColor("#9999bb"), 1.2))
                line.setZValue(2)
                self.addItem(line)
                self.aff_lines.append(line)

        # hooks: group by unordered pair
        pairs: Dict[Tuple[str, str], List[Hook]] = {}
        for h in self.project.hooks.values():
            if h.meta.is_deleted:
                continue
            a, b = h.a_character_id, h.b_character_id
            if a not in self.char_items or b not in self.char_items or a == b:
                continue
            key = tuple(sorted((a, b)))
            pairs.setdefault(key, []).append(h)

        for key, hooks in pairs.items():
            a, b = key
            a_it = self.char_items[a]
            b_it = self.char_items[b]
            p1 = a_it.pos()
            p2 = b_it.pos()

            n = len(hooks)
            hooks_sorted = sorted(hooks, key=lambda x: (x.type, x.label.casefold(), x.id))
            hovered = key in self.pair_hover
            spread = 26.0 if hovered else 12.0

            for i, h in enumerate(hooks_sorted):
                idx = i - (n - 1) / 2.0
                mid = (p1 + p2) / 2.0
                dx = p2.x() - p1.x()
                dy = p2.y() - p1.y()
                length = max((dx*dx + dy*dy) ** 0.5, 1.0)
                nx = -dy / length
                ny = dx / length
                ctrl = QtCore.QPointF(mid.x() + nx * idx * spread, mid.y() + ny * idx * spread)

                path = QtGui.QPainterPath(p1)
                path.quadTo(ctrl, p2)

                item = HookLineItem(h.id, h.label, h.type)
                item.pair_key = key
                item.setPath(path)
                self.addItem(item)
                self.addItem(item.text)
                self.hook_items[h.id] = item
                item.set_label_pos(ctrl.x() - 20, ctrl.y() - 10)

    def contextMenuEvent(self, event: QtWidgets.QGraphicsSceneContextMenuEvent):
        item = self.itemAt(event.scenePos(), QtGui.QTransform())

        # ПКМ по персонажу
        if isinstance(item, CharacterNodeItem):
            menu = QtWidgets.QMenu()
            act_edit = menu.addAction("Редактировать персонажа")
            act_hooks = menu.addAction("Завязки")
            act_goals = menu.addAction("Цели")
            menu.addSeparator()
            act_del = menu.addAction("Удалить персонажа")
            chosen = menu.exec(event.screenPos())
            if chosen == act_edit and self.on_open_character:
                self.on_open_character(item.char_id)
            elif chosen == act_hooks and self.on_open_hooks:
                self.on_open_hooks(item.char_id)
            elif chosen == act_goals and self.on_open_goals:
                self.on_open_goals(item.char_id)
            elif chosen == act_del and self.on_delete_character:
                self.on_delete_character(item.char_id)
            return

        # ПКМ по фракции
        if isinstance(item, FactionContainerItem):
            menu = QtWidgets.QMenu()
            act_del_f = menu.addAction("Удалить фракцию")
            chosen = menu.exec(event.screenPos())
            if chosen == act_del_f and self.on_delete_faction:
                self.on_delete_faction(item.faction_id)
            return

        # ПКМ по пустому полю: добавить персонажа / фракцию
        menu = QtWidgets.QMenu()
        act_add_char = menu.addAction("Добавить персонажа")
        act_add_fac = menu.addAction("Добавить фракцию")
        chosen = menu.exec(event.screenPos())
        if chosen == act_add_char and self.on_create_character_at:
            self.on_create_character_at(event.scenePos())
        elif chosen == act_add_fac and self.on_create_faction_at:
            self.on_create_faction_at(event.scenePos())

    def mouseDoubleClickEvent(self, event: QtWidgets.QGraphicsSceneMouseEvent):
        item = self.itemAt(event.scenePos(), QtGui.QTransform())
        if isinstance(item, CharacterNodeItem):
            if self.on_open_character:
                self.on_open_character(item.char_id)
            return
        if isinstance(item, FactionContainerItem):
            if self.on_open_faction:
                self.on_open_faction(item.faction_id)
            return
        super().mouseDoubleClickEvent(event)

    def hook_clicked(self, hook_id: str):

        if self.on_open_hook_by_id:
            self.on_open_hook_by_id(hook_id)

    def auto_layout(self):
        self._suppress = True
        import math
        factions = self.project.alive_factions()
        chars = self.project.alive_characters()

        r1 = 250
        for i, f in enumerate(factions):
            angle = (i / max(len(factions), 1)) * 2 * math.pi
            f.visual.x = r1 * math.cos(angle)
            f.visual.y = r1 * math.sin(angle)

        r2 = 430
        for i, c in enumerate(chars):
            angle = (i / max(len(chars), 1)) * 2 * math.pi
            c.visual.x = r2 * math.cos(angle)
            c.visual.y = r2 * math.sin(angle)

        for f in factions:
            it = self.faction_items.get(f.id)
            if it:
                it.setPos(f.visual.x, f.visual.y)
        for c in chars:
            it = self.char_items.get(c.id)
            if it:
                it.setPos(c.visual.x, c.visual.y)

        self._suppress = False

        self.selection_enabled = False
        self.draw_all_lines()


class FieldView(QtWidgets.QGraphicsView):
    zoomChanged = QtCore.Signal(int)  # percent
    def __init__(self, scene: FieldScene, parent=None):
        super().__init__(scene, parent)
        self.setRenderHints(QtGui.QPainter.Antialiasing | QtGui.QPainter.TextAntialiasing)
        self.setDragMode(QtWidgets.QGraphicsView.NoDrag)
        self.setTransformationAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.setResizeAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.setBackgroundBrush(QtGui.QBrush(QtGui.QColor("#0f1115")))
        self.setFrameShape(QtWidgets.QFrame.NoFrame)
        self.setFocusPolicy(QtCore.Qt.StrongFocus)

        # Панорамирование сцены ЛКМ по пустому месту (не ломая перетаскивание самих нод)
        self._panning = False
        self._pan_last = QtCore.QPoint()

        self._select_mode = False

    def current_zoom_percent(self) -> int:
        try:
            return int(round(self.transform().m11() * 100))
        except Exception:
            return 100

    def set_zoom_percent(self, percent: int):
        percent = max(10, min(400, int(percent)))
        self.resetTransform()
        f = percent / 100.0
        self.scale(f, f)
        try:
            self.zoomChanged.emit(self.current_zoom_percent())
        except Exception:
            pass

    def wheelEvent(self, event: QtGui.QWheelEvent):

        if event.modifiers() & QtCore.Qt.ControlModifier:
            factor = 1.15 if event.angleDelta().y() > 0 else 1/1.15
            self.scale(factor, factor)
            try:
                self.zoomChanged.emit(self.current_zoom_percent())
            except Exception:
                pass
        else:
            super().wheelEvent(event)

    def mousePressEvent(self, event: QtGui.QMouseEvent):
        if event.button() == QtCore.Qt.LeftButton:
            if not (event.modifiers() & QtCore.Qt.ControlModifier):
                try:
                    self.scene().clearSelection()
                except Exception:
                    pass
            # Панорамируем только если клик по пустому месту, чтобы ноды продолжали нормально двигаться
            scene_pos = self.mapToScene(event.pos())
            item = self.scene().itemAt(scene_pos, self.transform())
            if item is None:
                self._panning = True
                self._pan_last = event.pos()
                self.setCursor(QtCore.Qt.ClosedHandCursor)
                # чтобы не рисовалась рамка выделения
                self.setDragMode(QtWidgets.QGraphicsView.NoDrag)
                event.accept()
                return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event: QtGui.QMouseEvent):
        if self._panning:
            delta = event.pos() - self._pan_last
            self._pan_last = event.pos()
            # двигаем скроллы в противоположную сторону движению мыши
            h = self.horizontalScrollBar()
            v = self.verticalScrollBar()
            h.setValue(h.value() - delta.x())
            v.setValue(v.value() - delta.y())
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event: QtGui.QMouseEvent):
        if event.button() == QtCore.Qt.LeftButton and self._panning:
            self._panning = False
            self.unsetCursor()
            self.setDragMode(QtWidgets.QGraphicsView.NoDrag)
            event.accept()
            return
        super().mouseReleaseEvent(event)

    def _set_select_mode(self, enabled: bool):
        enabled = bool(enabled)
        if self._select_mode == enabled:
            return
        self._select_mode = enabled
        try:
            self.scene().set_selection_enabled(enabled)
        except Exception:
            pass
        self.setDragMode(QtWidgets.QGraphicsView.RubberBandDrag if enabled else QtWidgets.QGraphicsView.NoDrag)

    def _delete_selected(self):
        try:
            items = list(self.scene().selectedItems())
        except Exception:
            items = []
        if not items:
            return
        # map selected graphics items to character/faction ids
        char_ids = set()
        faction_ids = set()
        for it in items:
            cur = it
            # text items are usually children; walk up
            for _ in range(3):
                if cur is None:
                    break
                if isinstance(cur, CharacterNodeItem):
                    char_ids.add(cur.char_id)
                    break
                if isinstance(cur, FactionContainerItem):
                    faction_ids.add(cur.faction_id)
                    break
                cur = cur.parentItem()
        sc = self.scene()
        # delete chars first, then factions
        for cid in list(char_ids):
            cb = getattr(sc, "on_delete_character", None)
            if cb:
                try:
                    cb(cid)
                except Exception:
                    pass
        for fid in list(faction_ids):
            cb = getattr(sc, "on_delete_faction", None)
            if cb:
                try:
                    cb(fid)
                except Exception:
                    pass
        # avoid stale selection
        try:
            sc.clearSelection()
        except Exception:
            pass

    def keyPressEvent(self, event: QtGui.QKeyEvent):
        if event.key() == QtCore.Qt.Key_Control:
            self._set_select_mode(True)
            event.accept()
            return
        if event.key() == QtCore.Qt.Key_Delete:
            # Delete works on selected items (selection mode is enabled via Ctrl)
            self._delete_selected()
            event.accept()
            return
        super().keyPressEvent(event)

    def keyReleaseEvent(self, event: QtGui.QKeyEvent):
        if event.key() == QtCore.Qt.Key_Control:
            self._set_select_mode(False)
            event.accept()
            return
        super().keyReleaseEvent(event)




# =========================
# PDF Export
# =========================

def try_register_dejavu():
    candidates = [
        resource_path("DejaVuSans.ttf"),
        os.path.join(os.getcwd(), "DejaVuSans.ttf"),
        r"C:\Windows\Fonts\DejaVuSans.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont("AppFont", p))
                return "AppFont"
            except Exception:
                continue
    return None

def _resolve_pdf_font(font_name: str) -> str:
    std = {"Helvetica", "Times-Roman", "Courier"}
    if font_name in std:
        return font_name
    return try_register_dejavu() or "Helvetica"

def pdf_export_character(project: Project, char: Character, fields: dict, out_path: str, font_name: str = ""):
    font_name = _resolve_pdf_font(font_name or project.export_font)
    c = canvas.Canvas(out_path, pagesize=A4)
    w, h = A4
    x = 48
    y = h - 52

    def set_font(sz: int):
        c.setFont(font_name or "Helvetica", sz)

    def draw_line(text: str, sz=11, gap=14):
        nonlocal y
        text = strip_link_markers(text or "")
        set_font(sz)
        maxw = w - 2*x
        words = (text or "").split(" ")
        line = ""
        for ww in words:
            test = (line + " " + ww).strip()
            if c.stringWidth(test, font_name or "Helvetica", sz) > maxw and line:
                c.drawString(x, y, line)
                y -= gap
                line = ww
            else:
                line = test
        if line:
            c.drawString(x, y, line)
            y -= gap

        if y < 72:
            c.showPage()
            y = h - 52

    def header(txt: str):
        nonlocal y
        set_font(16)
        c.drawString(x, y, strip_link_markers(txt))
        y -= 22

    def section(txt: str):
        nonlocal y
        y -= 6
        set_font(13)
        c.drawString(x, y, strip_link_markers(txt))
        y -= 18

    if fields.get("name", True):
        header(char.name)

    if fields.get("factions", False):
        section("Фракции")
        names = []
        for fid in char.factions:
            f = project.factions.get(fid)
            if f and not f.meta.is_deleted:
                names.append(f.name)
        draw_line(", ".join(names) if names else "—")

    if fields.get("locations", False):
        section("Локации")
        names = []
        for lid in char.locations:
            l = project.locations.get(lid)
            if l and not l.meta.is_deleted:
                names.append(l.name)
        draw_line(", ".join(names) if names else "—")

    if fields.get("masters", False):
        section("Ответственный мастер")
        draw_line(", ".join(char.masters) if char.masters else "—")

    custom_mask = fields.get("custom", [False]*5)
    for i in range(5):
        if i < len(custom_mask) and custom_mask[i]:
            ttl = char.custom_tag_fields[i].field_title or f"Поле {i+1}"
            section(ttl)
            draw_line(", ".join(char.custom_tag_fields[i].values) if char.custom_tag_fields[i].values else "—")

    if fields.get("public", False):
        section("Публичный сюжет")
        txt = strip_link_markers(char.story_public or "")
        for ln in txt.splitlines() or ["—"]:
            draw_line(ln)

    if fields.get("private", False):
        section("Личный сюжет")
        txt = strip_link_markers(char.story_private or "")
        for ln in txt.splitlines() or ["—"]:
            draw_line(ln)

    # Goals: no status label, only title + description
    if fields.get("goals", False):
        section("Цели")
        gs = project.goals_for_character(char.id)
        if not gs:
            draw_line("—")
        else:
            for g in gs:
                draw_line(f"• {g.title}", sz=11)
                desc = strip_link_markers(g.description or "").strip()
                if desc:
                    for ln in desc.splitlines():
                        draw_line(f"  {ln}", sz=10, gap=13)

    # Hooks: only other name + description (no type/label)
    if fields.get("hooks", False):
        section("Завязки")
        hs = project.hooks_for_character(char.id)
        if not hs:
            draw_line("—")
        else:
            for hhk in hs:
                other_id = hhk.b_character_id if hhk.a_character_id == char.id else hhk.a_character_id
                other = project.characters.get(other_id)
                other_name = other.name if other and not other.meta.is_deleted else "???"
                draw_line(f"• {other_name}", sz=11)
                desc = strip_link_markers(hhk.description or "").strip()
                if desc:
                    for ln in desc.splitlines():
                        draw_line(f"  {ln}", sz=10, gap=13)
                else:
                    draw_line("  —", sz=10, gap=13)

    c.save()


# =========================
# DOCX Export
# =========================

def docx_export_character(project: Project, char: Character, fields: dict, out_path: str, font_name: str = ""):
    if Document is None:
        raise RuntimeError("python-docx not installed")

    doc = Document()
    font_name = (font_name or project.export_font or "").strip()
    if font_name:
        try:
            doc.styles["Normal"].font.name = font_name
        except Exception:
            pass

    def add_section(title: str):
        doc.add_heading(strip_link_markers(title), level=2)

    def add_line(text: str):
        doc.add_paragraph(strip_link_markers(text or "—"))

    if fields.get("name", True):
        doc.add_heading(strip_link_markers(char.name), level=1)

    if fields.get("factions", False):
        add_section("Фракции")
        names = []
        for fid in char.factions:
            f = project.factions.get(fid)
            if f and not f.meta.is_deleted:
                names.append(f.name)
        add_line(", ".join(names) if names else "—")

    if fields.get("locations", False):
        add_section("Локации")
        names = []
        for lid in char.locations:
            l = project.locations.get(lid)
            if l and not l.meta.is_deleted:
                names.append(l.name)
        add_line(", ".join(names) if names else "—")

    if fields.get("masters", False):
        add_section("Ответственный мастер")
        add_line(", ".join(char.masters) if char.masters else "—")

    custom_mask = fields.get("custom", [False] * 5)
    for i in range(5):
        if i < len(custom_mask) and custom_mask[i]:
            ttl = char.custom_tag_fields[i].field_title or f"Поле {i+1}"
            add_section(ttl)
            vals = char.custom_tag_fields[i].values
            add_line(", ".join(vals) if vals else "—")

    if fields.get("public", False):
        add_section("Публичный сюжет")
        txt = strip_link_markers(char.story_public or "")
        for ln in (txt.splitlines() or ["—"]):
            add_line(ln)

    if fields.get("private", False):
        add_section("Личный сюжет")
        txt = strip_link_markers(char.story_private or "")
        for ln in (txt.splitlines() or ["—"]):
            add_line(ln)

    if fields.get("goals", False):
        add_section("Цели")
        gs = project.goals_for_character(char.id)
        if not gs:
            add_line("—")
        else:
            for g in gs:
                doc.add_paragraph(f"• {g.title}")
                desc = strip_link_markers(g.description or "").strip()
                if desc:
                    for ln in desc.splitlines():
                        doc.add_paragraph(f"  {ln}")

    if fields.get("hooks", False):
        add_section("Завязки")
        hs = project.hooks_for_character(char.id)
        if not hs:
            add_line("—")
        else:
            for hhk in hs:
                other_id = hhk.b_character_id if hhk.a_character_id == char.id else hhk.a_character_id
                other = project.characters.get(other_id)
                other_name = other.name if other and not other.meta.is_deleted else "???"
                doc.add_paragraph(f"• {other_name}")
                desc = strip_link_markers(hhk.description or "").strip()
                if desc:
                    for ln in desc.splitlines():
                        doc.add_paragraph(f"  {ln}")
                else:
                    doc.add_paragraph("  —")

    doc.save(out_path)

# =========================
# Dialogs
# =========================

class ExportFieldsDialog(QtWidgets.QDialog):
    def __init__(self, title_map: List[str], defaults: Optional[dict] = None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Поля экспорта")
        self.resize(520, 560)
        layout = QtWidgets.QVBoxLayout(self)

        self.chk_name = QtWidgets.QCheckBox("Имя")
        self.chk_factions = QtWidgets.QCheckBox("Фракции")
        self.chk_locations = QtWidgets.QCheckBox("Локации")
        self.chk_masters = QtWidgets.QCheckBox("Ответственные мастера")
        self.chk_public = QtWidgets.QCheckBox("Публичный сюжет")
        self.chk_private = QtWidgets.QCheckBox("Личный сюжет")

        self.chk_goals = QtWidgets.QCheckBox("Цели персонажа")
        self.chk_hooks = QtWidgets.QCheckBox("Завязки персонажа")

        # defaults
        d = defaults or {}
        self.chk_name.setChecked(bool(d.get("name", True)))
        self.chk_factions.setChecked(bool(d.get("factions", True)))
        self.chk_locations.setChecked(bool(d.get("locations", True)))
        self.chk_masters.setChecked(bool(d.get("masters", False)))
        self.chk_public.setChecked(bool(d.get("public", False)))
        self.chk_private.setChecked(bool(d.get("private", False)))
        self.chk_goals.setChecked(bool(d.get("goals", False)))
        self.chk_hooks.setChecked(bool(d.get("hooks", False)))

        layout.addWidget(self.chk_name)
        layout.addWidget(self.chk_factions)
        layout.addWidget(self.chk_locations)
        layout.addWidget(self.chk_masters)

        layout.addSpacing(8)
        layout.addWidget(QtWidgets.QLabel("Пользовательские теги:"))
        self.custom_checks = []
        custom_defaults = d.get("custom") if isinstance(d.get("custom"), list) else [False] * 5
        for i in range(5):
            nm = title_map[i] if (i < len(title_map) and title_map[i]) else f"Поле {i+1}"
            chk = QtWidgets.QCheckBox(nm)
            if i < len(custom_defaults):
                chk.setChecked(bool(custom_defaults[i]))
            self.custom_checks.append(chk)
            layout.addWidget(chk)

        layout.addSpacing(8)
        layout.addWidget(self.chk_public)
        layout.addWidget(self.chk_private)

        layout.addSpacing(8)
        layout.addWidget(self.chk_goals)
        layout.addWidget(self.chk_hooks)

        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def get_fields(self) -> dict:
        return {
            "name": self.chk_name.isChecked(),
            "factions": self.chk_factions.isChecked(),
            "locations": self.chk_locations.isChecked(),
            "masters": self.chk_masters.isChecked(),
            "custom": [c.isChecked() for c in self.custom_checks],
            "public": self.chk_public.isChecked(),
            "private": self.chk_private.isChecked(),
            "goals": self.chk_goals.isChecked(),
            "hooks": self.chk_hooks.isChecked(),
        }


class SettingsTab(QtWidgets.QWidget):
    def __init__(self, project: Project, change_cb: Optional[Callable[[], None]] = None, parent=None):
        super().__init__(parent)
        self.project = project
        self._change_cb = change_cb

        root = QtWidgets.QVBoxLayout(self)
        root.setContentsMargins(16, 16, 16, 16)
        root.setSpacing(12)

        font_box = QtWidgets.QGroupBox("Экспорт: шрифт")
        fb = QtWidgets.QVBoxLayout(font_box)
        self.cmb_font = QtWidgets.QComboBox()
        # show all installed fonts
        try:
            families = QtGui.QFontDatabase().families()
        except Exception:
            families = []
        families = sorted({f for f in families if f})
        if "DejaVu Sans" not in families:
            families.insert(0, "DejaVu Sans")
        for f in families:
            self.cmb_font.addItem(f)
        fb.addWidget(self.cmb_font)
        root.addWidget(font_box)

        export_box = QtWidgets.QGroupBox("Экспорт: поля по умолчанию")
        eb = QtWidgets.QVBoxLayout(export_box)
        self.chk_name = QtWidgets.QCheckBox("Имя")
        self.chk_factions = QtWidgets.QCheckBox("Фракции")
        self.chk_locations = QtWidgets.QCheckBox("Локации")
        self.chk_masters = QtWidgets.QCheckBox("Ответственные мастера")
        self.chk_public = QtWidgets.QCheckBox("Публичный сюжет")
        self.chk_private = QtWidgets.QCheckBox("Личный сюжет")
        self.chk_goals = QtWidgets.QCheckBox("Цели персонажа")
        self.chk_hooks = QtWidgets.QCheckBox("Завязки персонажа")

        eb.addWidget(self.chk_name)
        eb.addWidget(self.chk_factions)
        eb.addWidget(self.chk_locations)
        eb.addWidget(self.chk_masters)

        eb.addSpacing(8)
        eb.addWidget(QtWidgets.QLabel("Пользовательские теги:"))
        self.custom_checks = []
        for i in range(5):
            chk = QtWidgets.QCheckBox(f"Поле {i+1}")
            self.custom_checks.append(chk)
            eb.addWidget(chk)

        eb.addSpacing(8)
        eb.addWidget(self.chk_public)
        eb.addWidget(self.chk_private)
        eb.addSpacing(8)
        eb.addWidget(self.chk_goals)
        eb.addWidget(self.chk_hooks)

        root.addWidget(export_box)
        root.addStretch(1)

        self.cmb_font.currentTextChanged.connect(self._on_change)
        for w in [
            self.chk_name, self.chk_factions, self.chk_locations, self.chk_masters,
            self.chk_public, self.chk_private, self.chk_goals, self.chk_hooks,
            *self.custom_checks,
        ]:
            w.stateChanged.connect(self._on_change)

        self.load_from_project()

    def load_from_project(self):
        # update labels for custom tags
        try:
            titles = list(self.project.character_custom_titles)
        except Exception:
            titles = [""] * 5
        for i, chk in enumerate(self.custom_checks):
            ttl = titles[i] if i < len(titles) and titles[i] else f"Поле {i+1}"
            chk.setText(ttl)

        # font
        font = (self.project.export_font or "").strip()
        idx = self.cmb_font.findText(font)
        if idx >= 0:
            self.cmb_font.setCurrentIndex(idx)
        else:
            if font:
                self.cmb_font.addItem(font)
                self.cmb_font.setCurrentText(font)

        d = self.project.export_defaults or {}
        self.chk_name.setChecked(bool(d.get("name", True)))
        self.chk_factions.setChecked(bool(d.get("factions", True)))
        self.chk_locations.setChecked(bool(d.get("locations", True)))
        self.chk_masters.setChecked(bool(d.get("masters", False)))
        self.chk_public.setChecked(bool(d.get("public", False)))
        self.chk_private.setChecked(bool(d.get("private", False)))
        self.chk_goals.setChecked(bool(d.get("goals", False)))
        self.chk_hooks.setChecked(bool(d.get("hooks", False)))
        custom = d.get("custom") if isinstance(d.get("custom"), list) else [False] * 5
        for i, chk in enumerate(self.custom_checks):
            chk.setChecked(bool(custom[i]) if i < len(custom) else False)

    def _on_change(self):
        self.project.export_font = self.cmb_font.currentText().strip() or "DejaVu Sans"
        self.project.export_defaults = {
            "name": self.chk_name.isChecked(),
            "factions": self.chk_factions.isChecked(),
            "locations": self.chk_locations.isChecked(),
            "masters": self.chk_masters.isChecked(),
            "custom": [c.isChecked() for c in self.custom_checks],
            "public": self.chk_public.isChecked(),
            "private": self.chk_private.isChecked(),
            "goals": self.chk_goals.isChecked(),
            "hooks": self.chk_hooks.isChecked(),
        }
        if self._change_cb:
            try:
                self._change_cb()
            except Exception:
                pass



class ExportSettingsDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, change_cb: Optional[Callable[[], None]] = None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("????????? ????????")
        self.resize(520, 640)

        root = QtWidgets.QVBoxLayout(self)
        self.settings = SettingsTab(project, change_cb=change_cb, parent=self)
        root.addWidget(self.settings, 1)

        btn_row = QtWidgets.QHBoxLayout()
        btn_row.addStretch(1)
        btn_close = QtWidgets.QPushButton("???????")
        btn_close.clicked.connect(self.accept)
        btn_row.addWidget(btn_close)
        root.addLayout(btn_row)
class GoalsDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, focus_char_id: str, open_by_name_cb: Callable[[str], None], parent=None):
        super().__init__(parent)
        self.project = project
        self.focus_char_id = focus_char_id
        self.open_by_name_cb = open_by_name_cb
        self.setWindowTitle("Цели персонажа")
        self.resize(980, 620)

        root = QtWidgets.QHBoxLayout(self)

        left = QtWidgets.QVBoxLayout()
        root.addLayout(left, 2)

        self.tbl = QtWidgets.QTableWidget(0, 2)
        self.tbl.setHorizontalHeaderLabels(["Статус", "Цель"])
        self.tbl.horizontalHeader().setStretchLastSection(True)
        self.tbl.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tbl.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        left.addWidget(self.tbl, 1)

        create_box = QtWidgets.QGroupBox("Создать цель")
        form = QtWidgets.QFormLayout(create_box)

        self.cmb_status_new = QtWidgets.QComboBox()
        self.cmb_status_new.addItems(["active", "completed"])
        self.ed_title_new = QtWidgets.QLineEdit()
        self.ed_title_new.setPlaceholderText("Короткая формулировка цели*")
        self.tx_desc_new = QtWidgets.QTextEdit()

        form.addRow("Статус", self.cmb_status_new)
        form.addRow("Название*", self.ed_title_new)
        form.addRow("Описание", self.tx_desc_new)

        self.btn_create = QtWidgets.QPushButton("Создать")
        form.addRow(self.btn_create)
        left.addWidget(create_box, 1)

        right = QtWidgets.QVBoxLayout()
        root.addLayout(right, 3)

        self.lbl_sel = QtWidgets.QLabel("Выбрана: —")
        right.addWidget(self.lbl_sel)

        self.cmb_status = QtWidgets.QComboBox()
        self.cmb_status.addItems(["active", "completed"])
        self.ed_title = QtWidgets.QLineEdit()

        def resolver(name: str):
            return self.project.get_object_by_name(name)

        self.tx_desc = HyperTextEdit(resolver)
        self.tx_desc.linkActivated.connect(self.open_by_name_cb)

        edit_form = QtWidgets.QFormLayout()
        edit_form.addRow("Статус", self.cmb_status)
        edit_form.addRow("Название*", self.ed_title)
        edit_form.addRow("Описание", self.tx_desc)
        right.addLayout(edit_form, 1)

        btns = QtWidgets.QHBoxLayout()
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        btns.addWidget(self.btn_delete)
        btns.addStretch(1)
        btns.addWidget(self.btn_save)
        right.addLayout(btns)

        self.selected_goal_id: Optional[str] = None

        self.btn_create.clicked.connect(self.create_goal)
        self.tbl.itemSelectionChanged.connect(self.on_select_row)
        self.btn_save.clicked.connect(self.save_selected)
        self.btn_delete.clicked.connect(self.delete_selected)

        self.refresh_table()

    def _status_label(self, st: str) -> str:
        return "Активная" if st == "active" else "Завершена"

    def refresh_table(self):
        rows = []
        for g in self.project.goals_for_character(self.focus_char_id):
            rows.append((g.status, g.title, g.id))
        self.tbl.setRowCount(0)
        for st, title, gid in rows:
            r = self.tbl.rowCount()
            self.tbl.insertRow(r)
            it0 = QtWidgets.QTableWidgetItem(self._status_label(st))
            it1 = QtWidgets.QTableWidgetItem(title)
            it0.setData(QtCore.Qt.UserRole, gid)
            it1.setData(QtCore.Qt.UserRole, gid)
            self.tbl.setItem(r, 0, it0)
            self.tbl.setItem(r, 1, it1)

    def on_select_row(self):
        items = self.tbl.selectedItems()
        if not items:
            self.selected_goal_id = None
            self.lbl_sel.setText("Выбрана: —")
            self.cmb_status.setCurrentText("active")
            self.ed_title.setText("")
            self.tx_desc.setPlainText("")
            return
        gid = items[0].data(QtCore.Qt.UserRole)
        self.selected_goal_id = gid
        g = self.project.goals.get(gid)
        if not g or g.meta.is_deleted:
            return
        self.lbl_sel.setText(f"Выбрана: {g.title}")
        self.cmb_status.setCurrentText(g.status)
        self.ed_title.setText(g.title)
        self.tx_desc.setPlainText(g.description)

    def create_goal(self):
        title = norm_spaces(self.ed_title_new.text())
        if not title:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Название цели обязательно.")
            return
        g = Goal(
            character_id=self.focus_char_id,
            status=self.cmb_status_new.currentText(),
            title=title,
            description=self.tx_desc_new.toPlainText()
        )
        self.project.goals[g.id] = g
        self.ed_title_new.clear()
        self.tx_desc_new.clear()
        self.refresh_table()

    def save_selected(self):
        if not self.selected_goal_id:
            return
        g = self.project.goals.get(self.selected_goal_id)
        if not g or g.meta.is_deleted:
            return
        title = norm_spaces(self.ed_title.text())
        if not title:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Название цели обязательно.")
            return
        g.status = self.cmb_status.currentText()
        g.title = title
        g.description = self.tx_desc.toPlainText()
        g.meta.touch("local")
        self.refresh_table()

    def delete_selected(self):
        if not self.selected_goal_id:
            return
        g = self.project.goals.get(self.selected_goal_id)
        if not g or g.meta.is_deleted:
            return
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить цель? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            g.meta.is_deleted = True
            g.meta.touch("local")
            self.selected_goal_id = None
            self.refresh_table()


class CharacterDialog(QtWidgets.QDialog):
    def __init__(
        self,
        project: Project,
        char: Character,
        open_by_name_cb: Callable[[str], None],
        open_hooks_cb: Callable[[str], None],
        open_goals_cb: Callable[[str], None],
                change_cb: Optional[Callable[[str], None]] = None,
        parent=None
    ):
        super().__init__(parent)
        self.project = project
        self.char = char
        self.open_by_name_cb = open_by_name_cb
        self.open_hooks_cb = open_hooks_cb
        self.open_goals_cb = open_goals_cb
        self._change_cb = change_cb
        self._live_timer = QtCore.QTimer(self)
        self._live_timer.setSingleShot(True)
        self._live_timer.timeout.connect(self._do_live_save)
        self._live_suspend = False

        # Ensure custom field titles are synchronized across all characters/factions
        self.project.apply_custom_field_titles()

        self.setWindowTitle("Настройки персонажа")
        self.resize(1040, 760)

        root = QtWidgets.QVBoxLayout(self)

        form = QtWidgets.QFormLayout()
        self.ed_name = QtWidgets.QLineEdit(char.name)
        form.addRow("Имя*", self.ed_name)
        root.addLayout(form)

        grid = QtWidgets.QGridLayout()
        root.addLayout(grid, 1)

        self.lst_factions = CheckListWidget()
        self.lst_locations = CheckListWidget()
        self.lst_stories = CheckListWidget()
        self.lst_masters = CheckListWidget()

        grid.addWidget(QtWidgets.QLabel("Фракции"), 0, 0)
        grid.addWidget(self.lst_factions, 1, 0)
        grid.addWidget(QtWidgets.QLabel("Локации"), 0, 1)
        grid.addWidget(self.lst_locations, 1, 1)
        grid.addWidget(QtWidgets.QLabel("Сюжеты"), 0, 2)
        grid.addWidget(self.lst_stories, 1, 2)
        grid.addWidget(QtWidgets.QLabel("Ответственный мастер"), 0, 3)
        grid.addWidget(self.lst_masters, 1, 3)

        custom_box = QtWidgets.QGroupBox("Пользовательские теги (5 полей)")
        custom_layout = QtWidgets.QFormLayout(custom_box)
        self.custom_title_edits = []
        self.custom_values_edits = []
        for i in range(5):
            title = QtWidgets.QLineEdit(char.custom_tag_fields[i].field_title)
            values = QtWidgets.QLineEdit(", ".join(char.custom_tag_fields[i].values))
            values.setPlaceholderText("значения через запятую")
            self.custom_title_edits.append(title)
            self.custom_values_edits.append(values)
            custom_layout.addRow(f"Поле {i+1} (название)", title)
            custom_layout.addRow(f"Поле {i+1} (значения)", values)
        root.addWidget(custom_box)

        story_split = QtWidgets.QSplitter(QtCore.Qt.Horizontal)
        root.addWidget(story_split, 2)

        def resolver(name: str):
            return self.project.get_object_by_name(name)

        self.tx_public = HyperTextEdit(resolver)
        self.tx_private = HyperTextEdit(resolver)
        self.tx_public.setPlainText(char.story_public)
        self.tx_private.setPlainText(char.story_private)
        self.tx_public.linkActivated.connect(self.open_by_name_cb)
        self.tx_private.linkActivated.connect(self.open_by_name_cb)

        left = QtWidgets.QWidget()
        lyt = QtWidgets.QVBoxLayout(left)
        lyt.addWidget(QtWidgets.QLabel("Публичный сюжет"))
        lyt.addWidget(self.tx_public, 1)
        right = QtWidgets.QWidget()
        ryt = QtWidgets.QVBoxLayout(right)
        ryt.addWidget(QtWidgets.QLabel("Личный сюжет"))
        ryt.addWidget(self.tx_private, 1)
        story_split.addWidget(left)
        story_split.addWidget(right)
        story_split.setSizes([1, 1])

        btn_row = QtWidgets.QHBoxLayout()
        root.addLayout(btn_row)

        self.btn_hooks = QtWidgets.QPushButton("Завязки…")
        self.btn_goals = QtWidgets.QPushButton("Цели…")
        self.btn_export = QtWidgets.QPushButton("Экспорт PDF…")
        self.btn_export_docx = QtWidgets.QPushButton("Экспорт DOCX…")
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        self.btn_cancel = QtWidgets.QPushButton("Отмена")

        btn_row.addWidget(self.btn_hooks)
        btn_row.addWidget(self.btn_goals)
        btn_row.addWidget(self.btn_export)
        btn_row.addWidget(self.btn_export_docx)
        btn_row.addWidget(self.btn_delete)
        btn_row.addStretch(1)
        btn_row.addWidget(self.btn_cancel)
        btn_row.addWidget(self.btn_save)

        self.btn_cancel.clicked.connect(self.reject)
        self.btn_save.clicked.connect(self.on_save)
        self.btn_delete.clicked.connect(self.on_delete)
        self.btn_export.clicked.connect(self.export_pdf)
        self.btn_export_docx.clicked.connect(self.export_docx)
        self.btn_hooks.clicked.connect(lambda: self.open_hooks_cb(self.char.id))
        self.btn_goals.clicked.connect(lambda: self.open_goals_cb(self.char.id))

        self.refresh_lists()

        # Live-sync: apply changes while user edits (debounced), so online clients update immediately
        self.ed_name.textEdited.connect(self._schedule_live_save)
        self.ed_name.textChanged.connect(self._schedule_live_save)
        self.tx_public.textChanged.connect(self._schedule_live_save)
        self.tx_private.textChanged.connect(self._schedule_live_save)
        self.lst_factions.itemChanged.connect(lambda *_: self._schedule_live_save())
        self.lst_locations.itemChanged.connect(lambda *_: self._schedule_live_save())
        self.lst_stories.itemChanged.connect(lambda *_: self._schedule_live_save())
        self.lst_masters.itemChanged.connect(lambda *_: self._schedule_live_save())
        for ed in self.custom_title_edits:
            ed.textEdited.connect(self._schedule_live_save)
            ed.textChanged.connect(self._schedule_live_save)
        for ed in self.custom_values_edits:
            ed.textEdited.connect(self._schedule_live_save)
            ed.textChanged.connect(self._schedule_live_save)

    def refresh_lists(self):
        factions = [f.name for f in self.project.alive_factions()]
        locations = [l.name for l in self.project.alive_locations()]
        stories = [s.name for s in self.project.alive_stories()]
        masters = sorted(list(self.project.masters_dict), key=lambda s: s.casefold())

        self.faction_name_to_id = {f.name: f.id for f in self.project.alive_factions()}
        self.location_name_to_id = {l.name: l.id for l in self.project.alive_locations()}
        self.story_name_to_id = {s.name: s.id for s in self.project.alive_stories()}

        checked_factions = set()
        for fid in self.char.factions:
            f = self.project.factions.get(fid)
            if f and not f.meta.is_deleted:
                checked_factions.add(f.name)

        checked_locations = set()
        for lid in self.char.locations:
            l = self.project.locations.get(lid)
            if l and not l.meta.is_deleted:
                checked_locations.add(l.name)

        checked_stories = set()
        for sid in self.char.stories:
            s = self.project.stories.get(sid)
            if s and not s.meta.is_deleted:
                checked_stories.add(s.name)

        self.lst_factions.set_items(factions, checked_factions)
        self.lst_locations.set_items(locations, checked_locations)
        self.lst_stories.set_items(stories, checked_stories)
        self.lst_masters.set_items(masters, set(self.char.masters))

    def export_pdf(self):
        title_map = [self.custom_title_edits[i].text().strip() for i in range(5)]
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        default_name = safe_filename(self.char.name) + ".pdf"
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Экспорт PDF", default_name, "PDF (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"

        self.on_save(silent=True)
        pdf_export_character(self.project, self.char, fields, path, font_name=self.project.export_font)
        QtWidgets.QMessageBox.information(self, "Готово", "PDF экспортирован.")

    def export_docx(self):
        if Document is None:
            QtWidgets.QMessageBox.critical(self, "Ошибка", "Модуль python-docx не установлен.")
            return
        title_map = [self.custom_title_edits[i].text().strip() for i in range(5)]
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        default_name = safe_filename(self.char.name) + ".docx"
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Экспорт DOCX", default_name, "DOCX (*.docx)")
        if not path:
            return
        if not path.lower().endswith(".docx"):
            path += ".docx"

        self.on_save(silent=True)
        try:
            docx_export_character(self.project, self.char, fields, path, font_name=self.project.export_font)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать DOCX:\n{e}")
            return
        QtWidgets.QMessageBox.information(self, "Готово", "DOCX экспортирован.")

    def on_delete(self):
        # If the project was reloaded by remote sync while this dialog is open,
        # self.char may be a stale object. Re-bind by id.
        try:
            cur = self.project.characters.get(self.char.id)
            if cur is not None:
                self.char = cur
        except Exception:
            pass
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить персонажа? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            self.char.meta.is_deleted = True
            self.char.meta.touch("local")
            self.accept()

    def on_save(self, silent: bool = False):
        # Re-bind to current project object (project may be reloaded by CRDT while dialog is open)
        try:
            cur = self.project.characters.get(self.char.id)
            if cur is not None:
                self.char = cur
        except Exception:
            pass
        # Rebuild name→id maps to keep checklist edits valid after remote updates
        self.faction_name_to_id = {f.name: f.id for f in self.project.alive_factions()}
        self.location_name_to_id = {l.name: l.id for l in self.project.alive_locations()}
        self.story_name_to_id = {s.name: s.id for s in self.project.alive_stories()}
        name = norm_spaces(self.ed_name.text())
        err = self.project.ensure_unique_name("character", self.char.id, name)
        if err and not silent:
            QtWidgets.QMessageBox.warning(self, "Ошибка", err)
            return
        # In silent/live mode we still want other fields to sync even if the name is temporarily invalid
        if not err:
            self.char.name = name

        sel_factions = self.lst_factions.checked_items()
        self.char.factions = [self.faction_name_to_id[n] for n in sel_factions if n in self.faction_name_to_id]

        sel_locations = self.lst_locations.checked_items()
        self.char.locations = [self.location_name_to_id[n] for n in sel_locations if n in self.location_name_to_id]

        sel_stories = self.lst_stories.checked_items()
        self.char.stories = [self.story_name_to_id[n] for n in sel_stories if n in self.story_name_to_id]

        self.char.masters = [norm_spaces(s) for s in self.lst_masters.checked_items() if norm_spaces(s)]
        for m in self.char.masters:
            self.project.masters_dict.add(m)

        for i in range(5):
            title = norm_spaces(self.custom_title_edits[i].text())
            # Sync title project-wide for characters (separate namespace from factions)
            self.project.set_character_custom_title(i, title)
            raw = self.custom_values_edits[i].text()
            vals = [norm_spaces(v) for v in raw.split(",")]
            vals = [v for v in vals if v]
            self.char.custom_tag_fields[i].values = sorted(list(dict.fromkeys(vals)), key=lambda s: s.casefold())
            for v in self.char.custom_tag_fields[i].values:
                self.project.character_custom_dicts[i].add(v)

        self.char.story_public = self.tx_public.toPlainText()
        self.char.story_private = self.tx_private.toPlainText()

        self.char.meta.touch("local")
        if not silent:
            self.accept()


    def _schedule_live_save(self):
        if getattr(self, '_live_suspend', False):
            return
        # Debounce to avoid flooding the server while typing
        self._live_timer.start(450)

    def _do_live_save(self):
        # Apply without closing dialog and notify main window for online sync
        try:
            self._live_suspend = True
            self.on_save(silent=True)
        finally:
            self._live_suspend = False
        cb = getattr(self, '_change_cb', None)
        if cb:
            try:
                cb(self.char.id)
            except Exception:
                pass

class StoryDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, story: Story, open_by_name_cb, parent=None):
        super().__init__(parent)
        self.project = project
        self.story = story
        self.open_by_name_cb = open_by_name_cb
        self.setWindowTitle("Сюжет")
        self.resize(760, 640)

        root = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.ed_name = QtWidgets.QLineEdit(story.name)
        form.addRow("Название*", self.ed_name)
        root.addLayout(form)

        def resolver(name: str):
            return self.project.get_object_by_name(name)

        self.tx_desc = HyperTextEdit(resolver)
        self.tx_desc.setPlainText(story.description)
        self.tx_desc.linkActivated.connect(self.open_by_name_cb)
        root.addWidget(QtWidgets.QLabel("Описание"))
        root.addWidget(self.tx_desc, 1)

        custom_box = QtWidgets.QGroupBox("Пользовательские теги сюжета (5 полей)")
        custom_layout = QtWidgets.QFormLayout(custom_box)
        self.custom_title_edits = []
        self.custom_values_edits = []
        for i in range(5):
            title = QtWidgets.QLineEdit(story.custom_tag_fields[i].field_title)
            values = QtWidgets.QLineEdit(", ".join(story.custom_tag_fields[i].values))
            values.setPlaceholderText("значения через запятую")
            self.custom_title_edits.append(title)
            self.custom_values_edits.append(values)
            custom_layout.addRow(f"Поле {i+1} (название)", title)
            custom_layout.addRow(f"Поле {i+1} (значения)", values)
        root.addWidget(custom_box)

        btn_row = QtWidgets.QHBoxLayout()
        root.addLayout(btn_row)
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        self.btn_cancel = QtWidgets.QPushButton("Отмена")
        btn_row.addWidget(self.btn_delete)
        btn_row.addStretch(1)
        btn_row.addWidget(self.btn_cancel)
        btn_row.addWidget(self.btn_save)

        self.btn_cancel.clicked.connect(self.reject)
        self.btn_save.clicked.connect(self.on_save)
        self.btn_delete.clicked.connect(self.on_delete)

    def on_delete(self):
        # Re-bind to current project object (project may be reloaded by CRDT while dialog is open)
        try:
            cur = self.project.factions.get(self.faction.id)
            if cur is not None:
                self.faction = cur
        except Exception:
            pass
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить сюжет? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            self.story.meta.is_deleted = True
            self.story.meta.touch("local")
            self.accept()

    def on_save(self):
        name = norm_spaces(self.ed_name.text())
        err = self.project.ensure_unique_name("story", self.story.id, name)
        if err:
            QtWidgets.QMessageBox.warning(self, "Ошибка", err)
            return
        self.story.name = name
        self.story.description = self.tx_desc.toPlainText()

        for i in range(5):
            self.story.custom_tag_fields[i].field_title = norm_spaces(self.custom_title_edits[i].text())
            raw = self.custom_values_edits[i].text()
            vals = [norm_spaces(v) for v in raw.split(",")]
            vals = [v for v in vals if v]
            self.story.custom_tag_fields[i].values = sorted(list(dict.fromkeys(vals)), key=lambda s: s.casefold())
            for v in self.story.custom_tag_fields[i].values:
                self.project.story_custom_dicts[i].add(v)

        self.story.meta.touch("local")
        self.accept()


class FactionDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, faction: Faction, open_by_name_cb, change_cb: Optional[Callable[[str], None]] = None, parent=None):
        super().__init__(parent)
        self.project = project
        self.faction = faction
        self.open_by_name_cb = open_by_name_cb
        self._change_cb = change_cb
        self._live_timer = QtCore.QTimer(self)
        self._live_timer.setSingleShot(True)
        self._live_timer.timeout.connect(self._do_live_save)
        self._live_suspend = False


        # Ensure custom field titles are synchronized across all characters/factions
        self.project.apply_custom_field_titles()
        self.setWindowTitle("Фракция")
        self.resize(760, 640)

        root = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.ed_name = QtWidgets.QLineEdit(faction.name)
        form.addRow("Название*", self.ed_name)
        root.addLayout(form)

        def resolver(name: str):
            return self.project.get_object_by_name(name)

        self.tx_desc = HyperTextEdit(resolver)
        self.tx_desc.setPlainText(faction.description)
        self.tx_desc.linkActivated.connect(self.open_by_name_cb)
        root.addWidget(QtWidgets.QLabel("Описание"))
        root.addWidget(self.tx_desc, 1)

        custom_box = QtWidgets.QGroupBox("Пользовательские теги фракции (5 полей)")
        custom_layout = QtWidgets.QFormLayout(custom_box)
        self.custom_title_edits = []
        self.custom_values_edits = []
        for i in range(5):
            title = QtWidgets.QLineEdit(faction.custom_tag_fields[i].field_title)
            values = QtWidgets.QLineEdit(", ".join(faction.custom_tag_fields[i].values))
            values.setPlaceholderText("значения через запятую")
            self.custom_title_edits.append(title)
            self.custom_values_edits.append(values)
            custom_layout.addRow(f"Поле {i+1} (название)", title)
            custom_layout.addRow(f"Поле {i+1} (значения)", values)
        root.addWidget(custom_box)

        btn_row = QtWidgets.QHBoxLayout()
        root.addLayout(btn_row)
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        self.btn_cancel = QtWidgets.QPushButton("Отмена")
        btn_row.addWidget(self.btn_delete)
        btn_row.addStretch(1)
        btn_row.addWidget(self.btn_cancel)
        btn_row.addWidget(self.btn_save)

        self.btn_cancel.clicked.connect(self.reject)
        self.btn_save.clicked.connect(self.on_save)
        self.btn_delete.clicked.connect(self.on_delete)

        # Live-sync: apply changes while user edits (debounced) so online clients see updates immediately
        self.ed_name.textEdited.connect(self._schedule_live_save)
        self.ed_name.textChanged.connect(self._schedule_live_save)
        self.tx_desc.textChanged.connect(self._schedule_live_save)
        for ed in self.custom_title_edits:
            ed.textEdited.connect(self._schedule_live_save)
            ed.textChanged.connect(self._schedule_live_save)
        for ed in self.custom_values_edits:
            ed.textEdited.connect(self._schedule_live_save)
            ed.textChanged.connect(self._schedule_live_save)

    def on_delete(self):
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить фракцию? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            self.faction.meta.is_deleted = True
            self.faction.meta.touch("local")
            self.accept()

    def on_save(self, silent: bool = False):
        # Re-bind to current project object (project may be reloaded by CRDT while dialog is open)
        try:
            cur = self.project.factions.get(self.faction.id)
            if cur is not None:
                self.faction = cur
        except Exception:
            pass
        name = norm_spaces(self.ed_name.text())
        err = self.project.ensure_unique_name("faction", self.faction.id, name)
        if err and not silent:
            QtWidgets.QMessageBox.warning(self, "Ошибка", err)
            return
        if not err:
            self.faction.name = name
        self.faction.description = self.tx_desc.toPlainText()

        for i in range(5):
            title = norm_spaces(self.custom_title_edits[i].text())
            # Sync title project-wide for factions (separate namespace from characters)
            self.project.set_faction_custom_title(i, title)
            raw = self.custom_values_edits[i].text()
            vals = [norm_spaces(v) for v in raw.split(",")]
            vals = [v for v in vals if v]
            self.faction.custom_tag_fields[i].values = sorted(list(dict.fromkeys(vals)), key=lambda s: s.casefold())
            for v in self.faction.custom_tag_fields[i].values:
                self.project.faction_custom_dicts[i].add(v)

        self.faction.meta.touch("local")
        if not silent:
            self.accept()

    def _schedule_live_save(self):
        if getattr(self, '_live_suspend', False):
            return
        self._live_timer.start(450)

    def _do_live_save(self):
        try:
            self._live_suspend = True
            self.on_save(silent=True)
        finally:
            self._live_suspend = False
        cb = getattr(self, '_change_cb', None)
        if cb:
            try:
                cb(self.faction.id)
            except Exception:
                pass


class LocationDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, location: Location, open_by_name_cb, parent=None):
        super().__init__(parent)
        self.project = project
        self.location = location
        self.open_by_name_cb = open_by_name_cb
        self.setWindowTitle("Локация")
        self.resize(760, 640)

        root = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.ed_name = QtWidgets.QLineEdit(location.name)
        form.addRow("Название*", self.ed_name)
        root.addLayout(form)

        def resolver(name: str):
            return self.project.get_object_by_name(name)

        self.tx_desc = HyperTextEdit(resolver)
        self.tx_desc.setPlainText(location.description)
        self.tx_desc.linkActivated.connect(self.open_by_name_cb)
        root.addWidget(QtWidgets.QLabel("Описание"))
        root.addWidget(self.tx_desc, 1)

        custom_box = QtWidgets.QGroupBox("Пользовательские теги локации (5 полей)")
        custom_layout = QtWidgets.QFormLayout(custom_box)
        self.custom_title_edits = []
        self.custom_values_edits = []
        for i in range(5):
            title = QtWidgets.QLineEdit(location.custom_tag_fields[i].field_title)
            values = QtWidgets.QLineEdit(", ".join(location.custom_tag_fields[i].values))
            values.setPlaceholderText("значения через запятую")
            self.custom_title_edits.append(title)
            self.custom_values_edits.append(values)
            custom_layout.addRow(f"Поле {i+1} (название)", title)
            custom_layout.addRow(f"Поле {i+1} (значения)", values)
        root.addWidget(custom_box)

        btn_row = QtWidgets.QHBoxLayout()
        root.addLayout(btn_row)
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        self.btn_cancel = QtWidgets.QPushButton("Отмена")
        btn_row.addWidget(self.btn_delete)
        btn_row.addStretch(1)
        btn_row.addWidget(self.btn_cancel)
        btn_row.addWidget(self.btn_save)

        self.btn_cancel.clicked.connect(self.reject)
        self.btn_save.clicked.connect(self.on_save)
        self.btn_delete.clicked.connect(self.on_delete)

    def on_delete(self):
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить локацию? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            self.location.meta.is_deleted = True
            self.location.meta.touch("local")
            self.accept()

    def on_save(self):
        name = norm_spaces(self.ed_name.text())
        err = self.project.ensure_unique_name("location", self.location.id, name)
        if err:
            QtWidgets.QMessageBox.warning(self, "Ошибка", err)
            return
        self.location.name = name
        self.location.description = self.tx_desc.toPlainText()

        for i in range(5):
            self.location.custom_tag_fields[i].field_title = norm_spaces(self.custom_title_edits[i].text())
            raw = self.custom_values_edits[i].text()
            vals = [norm_spaces(v) for v in raw.split(",")]
            vals = [v for v in vals if v]
            self.location.custom_tag_fields[i].values = sorted(list(dict.fromkeys(vals)), key=lambda s: s.casefold())
            for v in self.location.custom_tag_fields[i].values:
                self.project.location_custom_dicts[i].add(v)

        self.location.meta.touch("local")
        self.accept()


class HooksDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, focus_char_id: Optional[str] = None, focus_hook_id: Optional[str] = None, parent=None):
        super().__init__(parent)
        self.project = project
        self.focus_char_id = focus_char_id
        self.focus_hook_id = focus_hook_id
        self.setWindowTitle("Завязки")
        self.resize(980, 640)

        root = QtWidgets.QHBoxLayout(self)

        left = QtWidgets.QVBoxLayout()
        root.addLayout(left, 2)

        self.tbl = QtWidgets.QTableWidget(0, 4)
        self.tbl.setHorizontalHeaderLabels(["A", "B", "Тип", "Подпись"])
        self.tbl.horizontalHeader().setStretchLastSection(True)
        self.tbl.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tbl.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        left.addWidget(self.tbl, 1)

        create_box = QtWidgets.QGroupBox("Создать завязку")
        create_form = QtWidgets.QFormLayout(create_box)

        self.cmb_a = QtWidgets.QComboBox()
        self.cmb_b = QtWidgets.QComboBox()
        self.cmb_type = QtWidgets.QComboBox()
        self.cmb_type.addItems(["positive", "conflict"])
        self.ed_label = QtWidgets.QLineEdit()
        self.tx_desc = QtWidgets.QTextEdit()

        create_form.addRow("Персонаж A", self.cmb_a)
        create_form.addRow("Персонаж B", self.cmb_b)
        create_form.addRow("Тип", self.cmb_type)
        create_form.addRow("Подпись*", self.ed_label)
        create_form.addRow("Описание", self.tx_desc)

        self.btn_create = QtWidgets.QPushButton("Создать")
        create_form.addRow(self.btn_create)
        left.addWidget(create_box, 1)

        right = QtWidgets.QVBoxLayout()
        root.addLayout(right, 3)

        self.lbl_sel = QtWidgets.QLabel("Выбрана: —")
        right.addWidget(self.lbl_sel)

        self.sel_type = QtWidgets.QComboBox()
        self.sel_type.addItems(["positive", "conflict"])
        self.sel_label = QtWidgets.QLineEdit()
        self.sel_desc = QtWidgets.QTextEdit()

        form = QtWidgets.QFormLayout()
        form.addRow("Тип", self.sel_type)
        form.addRow("Подпись*", self.sel_label)
        form.addRow("Описание", self.sel_desc)
        right.addLayout(form, 1)

        btns = QtWidgets.QHBoxLayout()
        self.btn_save = QtWidgets.QPushButton("Сохранить")
        self.btn_delete = QtWidgets.QPushButton("Удалить")
        btns.addWidget(self.btn_delete)
        btns.addStretch(1)
        btns.addWidget(self.btn_save)
        right.addLayout(btns)

        self.selected_hook_id: Optional[str] = None

        self.btn_create.clicked.connect(self.create_hook)
        self.tbl.itemSelectionChanged.connect(self.on_select_row)
        self.btn_save.clicked.connect(self.save_selected)
        self.btn_delete.clicked.connect(self.delete_selected)

        self.populate_chars()
        self.refresh_table()

        if focus_hook_id:
            self.select_hook(focus_hook_id)

    def populate_chars(self):
        chars = self.project.alive_characters()
        self.char_name_by_id = {c.id: c.name for c in chars}
        names = [c.name for c in sorted(chars, key=lambda x: x.name.casefold())]
        self.cmb_a.clear()
        self.cmb_b.clear()
        self.cmb_a.addItems(names)
        self.cmb_b.addItems(names)

        if self.focus_char_id and self.focus_char_id in self.char_name_by_id:
            nm = self.char_name_by_id[self.focus_char_id]
            ia = self.cmb_a.findText(nm)
            if ia >= 0:
                self.cmb_a.setCurrentIndex(ia)

    def refresh_table(self):
        rows = []
        for h in self.project.hooks.values():
            if h.meta.is_deleted:
                continue
            if self.focus_char_id:
                if h.a_character_id != self.focus_char_id and h.b_character_id != self.focus_char_id:
                    continue
            a = self.char_name_by_id.get(h.a_character_id, "?")
            b = self.char_name_by_id.get(h.b_character_id, "?")
            rows.append((a, b, h.type, h.label, h.id))
        rows.sort(key=lambda x: (x[0].casefold(), x[1].casefold(), x[2], x[3].casefold()))

        self.tbl.setRowCount(0)
        for a, b, t, lbl, hid in rows:
            r = self.tbl.rowCount()
            self.tbl.insertRow(r)
            for col, val in enumerate([a, b, t, lbl]):
                it = QtWidgets.QTableWidgetItem(val)
                it.setData(QtCore.Qt.UserRole, hid)
                self.tbl.setItem(r, col, it)

    def select_hook(self, hook_id: str):
        for r in range(self.tbl.rowCount()):
            it = self.tbl.item(r, 0)
            if it and it.data(QtCore.Qt.UserRole) == hook_id:
                self.tbl.selectRow(r)
                return

    def on_select_row(self):
        items = self.tbl.selectedItems()
        if not items:
            self.selected_hook_id = None
            self.lbl_sel.setText("Выбрана: —")
            return
        hook_id = items[0].data(QtCore.Qt.UserRole)
        self.selected_hook_id = hook_id
        h = self.project.hooks.get(hook_id)
        if not h:
            return
        a = self.char_name_by_id.get(h.a_character_id, "?")
        b = self.char_name_by_id.get(h.b_character_id, "?")
        self.lbl_sel.setText(f"Выбрана: {a} — {b}")
        self.sel_type.setCurrentText(h.type)
        self.sel_label.setText(h.label)
        self.sel_desc.setPlainText(h.description)

    def _name_to_id(self, name: str) -> Optional[str]:
        for cid, nm in self.char_name_by_id.items():
            if nm == name:
                return cid
        return None

    def create_hook(self):
        a_name = self.cmb_a.currentText()
        b_name = self.cmb_b.currentText()
        a_id = self._name_to_id(a_name)
        b_id = self._name_to_id(b_name)
        if not a_id or not b_id or a_id == b_id:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Выбери двух разных персонажей.")
            return
        label = norm_spaces(self.ed_label.text())
        if not label:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Подпись обязательна.")
            return
        h = Hook(
            a_character_id=a_id,
            b_character_id=b_id,
            type=self.cmb_type.currentText(),
            label=label,
            description=self.tx_desc.toPlainText()
        )
        self.project.hooks[h.id] = h
        self.refresh_table()
        self.ed_label.clear()
        self.tx_desc.clear()
        self.select_hook(h.id)

    def save_selected(self):
        if not self.selected_hook_id:
            return
        h = self.project.hooks.get(self.selected_hook_id)
        if not h:
            return
        label = norm_spaces(self.sel_label.text())
        if not label:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Подпись обязательна.")
            return
        h.type = self.sel_type.currentText()
        h.label = label
        h.description = self.sel_desc.toPlainText()
        h.meta.touch("local")
        self.refresh_table()
        self.select_hook(h.id)

    def delete_selected(self):
        if not self.selected_hook_id:
            return
        h = self.project.hooks.get(self.selected_hook_id)
        if not h:
            return
        r = QtWidgets.QMessageBox.question(self, "Удаление", "Удалить завязку? (мягкое удаление)")
        if r == QtWidgets.QMessageBox.Yes:
            h.meta.is_deleted = True
            h.meta.touch("local")
            self.selected_hook_id = None
            self.refresh_table()


class HistoryDialog(QtWidgets.QDialog):
    def __init__(self, project: Project, parent=None):
        super().__init__(parent)
        self.project = project
        self.setWindowTitle("История версий")
        self.resize(820, 520)

        root = QtWidgets.QVBoxLayout(self)
        self.tbl = QtWidgets.QTableWidget(0, 4)
        self.tbl.setHorizontalHeaderLabels(["Время", "Тип", "Автор", "Комментарий"])
        self.tbl.horizontalHeader().setStretchLastSection(True)
        self.tbl.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tbl.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        root.addWidget(self.tbl, 1)

        btns = QtWidgets.QHBoxLayout()
        self.btn_restore = QtWidgets.QPushButton("Откатиться к версии")
        self.btn_close = QtWidgets.QPushButton("Закрыть")
        btns.addWidget(self.btn_restore)
        btns.addStretch(1)
        btns.addWidget(self.btn_close)
        root.addLayout(btns)

        self.btn_close.clicked.connect(self.reject)
        self.btn_restore.clicked.connect(self.restore)
        self.populate()

    def populate(self):
        self.tbl.setRowCount(0)
        for ve in reversed(self.project.versions):
            r = self.tbl.rowCount()
            self.tbl.insertRow(r)
            vals = [ve.timestamp, ve.kind, ve.author, ve.comment]
            for c, v in enumerate(vals):
                it = QtWidgets.QTableWidgetItem(v)
                it.setData(QtCore.Qt.UserRole, ve.id)
                self.tbl.setItem(r, c, it)

    def restore(self):
        items = self.tbl.selectedItems()
        if not items:
            return
        vid = items[0].data(QtCore.Qt.UserRole)
        ve = next((x for x in self.project.versions if x.id == vid), None)
        if not ve:
            return
        r = QtWidgets.QMessageBox.question(self, "Откат", "Откатить проект к выбранной версии?")
        if r != QtWidgets.QMessageBox.Yes:
            return
        self.project.load_state(ve.state)
        self.accept()


# =========================
# Tabs
# =========================

class CharactersTab(QtWidgets.QWidget):
    def __init__(self, project: Project, open_character_cb, change_cb=None, parent=None):
        super().__init__(parent)
        self.project = project
        self.open_character_cb = open_character_cb
        self.change_cb = change_cb

        root = QtWidgets.QVBoxLayout(self)

        top = QtWidgets.QHBoxLayout()
        root.addLayout(top)

        self.ed_search = QtWidgets.QLineEdit()
        self.ed_search.setPlaceholderText("Поиск по имени…")
        top.addWidget(self.ed_search, 1)

        self.cmb_logic = QtWidgets.QComboBox()
        self.cmb_logic.addItems(["AND", "OR"])
        top.addWidget(QtWidgets.QLabel("Фильтр логика:"))
        top.addWidget(self.cmb_logic)

        self.btn_new = QtWidgets.QPushButton("Создать персонажа")
        top.addWidget(self.btn_new)

        self.splitter = QtWidgets.QSplitter(QtCore.Qt.Vertical)
        root.addWidget(self.splitter, 1)

        self.filters_group = QtWidgets.QGroupBox("Фильтры по тегам")
        flt_l = QtWidgets.QGridLayout(self.filters_group)

        self.flt_factions = CheckListWidget()
        self.flt_locations = CheckListWidget()
        self.flt_stories = CheckListWidget()
        self.flt_masters = CheckListWidget()

        flt_l.addWidget(QtWidgets.QLabel("Фракции"), 0, 0)
        flt_l.addWidget(self.flt_factions, 1, 0)
        flt_l.addWidget(QtWidgets.QLabel("Локации"), 0, 1)
        flt_l.addWidget(self.flt_locations, 1, 1)
        flt_l.addWidget(QtWidgets.QLabel("Сюжеты"), 0, 2)
        flt_l.addWidget(self.flt_stories, 1, 2)
        flt_l.addWidget(QtWidgets.QLabel("Мастера"), 0, 3)
        flt_l.addWidget(self.flt_masters, 1, 3)

        self.flt_custom = []
        for i in range(5):
            w = CheckListWidget()
            self.flt_custom.append(w)
            r = 2 + (i // 4) * 2
            c = i % 4
            flt_l.addWidget(QtWidgets.QLabel(f"Кастом {i+1}"), r, c)
            flt_l.addWidget(w, r + 1, c)

        self.filters_scroll = QtWidgets.QScrollArea()
        self.filters_scroll.setWidgetResizable(True)
        self.filters_scroll.setFrameShape(QtWidgets.QFrame.NoFrame)
        self.filters_scroll.setWidget(self.filters_group)

        self.tbl = QtWidgets.QTableWidget(0, 5)
        self.tbl.setHorizontalHeaderLabels(["Имя", "Фракции", "Локации", "Сюжеты", "Мастера"])
        self.tbl.horizontalHeader().setStretchLastSection(True)
        self.tbl.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tbl.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)

        self.splitter.addWidget(self.filters_scroll)
        self.splitter.addWidget(self.tbl)
        self.splitter.setStretchFactor(0, 1)
        self.splitter.setStretchFactor(1, 3)
        self.splitter.setSizes([280, 520])

        self.ed_search.textChanged.connect(self.refresh)
        self.cmb_logic.currentTextChanged.connect(self.refresh)
        self.btn_new.clicked.connect(self.create_new)
        self.tbl.cellDoubleClicked.connect(self.open_row)

        for w in [self.flt_factions, self.flt_locations, self.flt_stories, self.flt_masters] + self.flt_custom:
            w.itemChanged.connect(self.refresh)

        self.refresh_filters()
        self.refresh()

    def get_selected_character_id(self) -> Optional[str]:
        sel = self.tbl.selectionModel().selectedRows()
        if not sel:
            return None
        row = sel[0].row()
        it = self.tbl.item(row, 0)
        if not it:
            return None
        return it.data(QtCore.Qt.UserRole)

    def get_selected_character_ids(self) -> List[str]:
        ids = []
        for idx in self.tbl.selectionModel().selectedRows():
            it = self.tbl.item(idx.row(), 0)
            if it:
                cid = it.data(QtCore.Qt.UserRole)
                if cid:
                    ids.append(cid)
        # unique preserve order
        seen = set()
        out = []
        for cid in ids:
            if cid not in seen:
                seen.add(cid)
                out.append(cid)
        return out

    def refresh_filters(self):
        prev_factions = set(self.flt_factions.checked_items())
        prev_locations = set(self.flt_locations.checked_items())
        prev_stories = set(self.flt_stories.checked_items())
        prev_masters = set(self.flt_masters.checked_items())
        prev_custom = [set(w.checked_items()) for w in self.flt_custom]

        factions = [f.name for f in self.project.alive_factions()]
        locations = [l.name for l in self.project.alive_locations()]
        stories = [s.name for s in self.project.alive_stories()]
        masters = sorted(list(self.project.masters_dict), key=lambda s: s.casefold())

        blockers = [
            QtCore.QSignalBlocker(self.flt_factions),
            QtCore.QSignalBlocker(self.flt_locations),
            QtCore.QSignalBlocker(self.flt_stories),
            QtCore.QSignalBlocker(self.flt_masters),
        ] + [QtCore.QSignalBlocker(w) for w in self.flt_custom]

        self.flt_factions.set_items(factions, prev_factions.intersection(set(factions)))
        self.flt_locations.set_items(locations, prev_locations.intersection(set(locations)))
        self.flt_stories.set_items(stories, prev_stories.intersection(set(stories)))
        self.flt_masters.set_items(masters, prev_masters.intersection(set(masters)))

        for i in range(5):
            items = sorted(list(self.project.character_custom_dicts[i]), key=lambda s: s.casefold())
            self.flt_custom[i].set_items(items, prev_custom[i].intersection(set(items)))

        _ = blockers

    def create_new(self):
        c = Character(name="Новый персонаж")
        c.visual.x, c.visual.y = 0.0, 0.0
        self.project.characters[c.id] = c
        if self.change_cb:
            self.change_cb()
        # Apply global custom field titles so the dialog shows consistent names
        self.project.apply_custom_field_titles()
        self.open_character_cb(c.id)
        self.refresh_filters()
        self.refresh()

    def open_row(self, row: int, col: int):
        it = self.tbl.item(row, 0)
        if not it:
            return
        cid = it.data(QtCore.Qt.UserRole)
        if cid:
            self.open_character_cb(cid)

    def _selected_filters(self):
        return {
            "factions": set(self.flt_factions.checked_items()),
            "locations": set(self.flt_locations.checked_items()),
            "stories": set(self.flt_stories.checked_items()),
            "masters": set(self.flt_masters.checked_items()),
            "custom": [set(w.checked_items()) for w in self.flt_custom],
        }

    def refresh(self):
        q = norm_key(self.ed_search.text())
        logic = self.cmb_logic.currentText()
        flt = self._selected_filters()

        faction_name_to_id = {f.name: f.id for f in self.project.alive_factions()}
        location_name_to_id = {l.name: l.id for l in self.project.alive_locations()}
        story_name_to_id = {s.name: s.id for s in self.project.alive_stories()}

        def char_matches(c: Character) -> bool:
            if c.meta.is_deleted:
                return False
            if q and q not in norm_key(c.name):
                return False

            conditions = []

            if flt["factions"]:
                selected_ids = {faction_name_to_id[n] for n in flt["factions"] if n in faction_name_to_id}
                conditions.append(bool(selected_ids.intersection(set(c.factions))))

            if flt["locations"]:
                selected_ids = {location_name_to_id[n] for n in flt["locations"] if n in location_name_to_id}
                conditions.append(bool(selected_ids.intersection(set(c.locations))))

            if flt["stories"]:
                selected_ids = {story_name_to_id[n] for n in flt["stories"] if n in story_name_to_id}
                conditions.append(bool(selected_ids.intersection(set(c.stories))))

            if flt["masters"]:
                conditions.append(bool(set(c.masters).intersection(flt["masters"])))

            for i in range(5):
                if flt["custom"][i]:
                    conditions.append(bool(set(c.custom_tag_fields[i].values).intersection(flt["custom"][i])))

            if not conditions:
                return True
            return all(conditions) if logic == "AND" else any(conditions)

        rows = []
        for c in self.project.alive_characters():
            if not char_matches(c):
                continue
            facs = []
            for fid in c.factions:
                f = self.project.factions.get(fid)
                if f and not f.meta.is_deleted:
                    facs.append(f.name)
            locs = []
            for lid in c.locations:
                l = self.project.locations.get(lid)
                if l and not l.meta.is_deleted:
                    locs.append(l.name)
            stys = []
            for sid in c.stories:
                s = self.project.stories.get(sid)
                if s and not s.meta.is_deleted:
                    stys.append(s.name)
            rows.append((c.name, ", ".join(facs), ", ".join(locs), ", ".join(stys), ", ".join(c.masters), c.id))

        rows.sort(key=lambda x: x[0].casefold())

        self.tbl.setRowCount(0)
        for nm, fac, loc, sty, mas, cid in rows:
            r = self.tbl.rowCount()
            self.tbl.insertRow(r)
            for col, val in enumerate([nm, fac, loc, sty, mas]):
                it = QtWidgets.QTableWidgetItem(val)
                it.setData(QtCore.Qt.UserRole, cid)
                self.tbl.setItem(r, col, it)


class SimpleListTab(QtWidgets.QWidget):
    def __init__(self, title: str, project: Project, kind: str, open_cb, create_cb, parent=None):
        super().__init__(parent)
        self.project = project
        self.kind = kind
        self.open_cb = open_cb
        self.create_cb = create_cb

        root = QtWidgets.QVBoxLayout(self)
        top = QtWidgets.QHBoxLayout()
        root.addLayout(top)

        self.ed_search = QtWidgets.QLineEdit()
        self.ed_search.setPlaceholderText("Поиск…")
        self.btn_new = QtWidgets.QPushButton(f"Создать {title.lower()}")
        top.addWidget(self.ed_search, 1)
        top.addWidget(self.btn_new)

        self.lst = QtWidgets.QListWidget()
        root.addWidget(self.lst, 1)

        self.btn_new.clicked.connect(self.create_cb)
        self.ed_search.textChanged.connect(self.refresh)
        self.lst.itemDoubleClicked.connect(lambda it: self.open_cb(it.data(QtCore.Qt.UserRole)))

        self.refresh()

    def refresh(self):
        q = norm_key(self.ed_search.text())
        self.lst.clear()
        if self.kind == "faction":
            items = [(f.name, f.id) for f in self.project.alive_factions() if (not q or q in norm_key(f.name))]
        elif self.kind == "location":
            items = [(l.name, l.id) for l in self.project.alive_locations() if (not q or q in norm_key(l.name))]
        else:  # story
            items = [(s.name, s.id) for s in self.project.alive_stories() if (not q or q in norm_key(s.name))]
        items.sort(key=lambda x: x[0].casefold())
        for name, oid in items:
            it = QtWidgets.QListWidgetItem(name)
            it.setData(QtCore.Qt.UserRole, oid)
            self.lst.addItem(it)

# =========================
# CRDT Sync (Y-Py)
# =========================

try:
    from y_py import YDoc, YMap, YArray, YText, apply_update as y_apply_update
    from y_py import encode_state_vector as y_encode_state_vector
    from y_py import encode_state_as_update as y_encode_state_as_update
except Exception:  # y-py is optional at runtime (offline/editor-only builds)
    YDoc = None
    YMap = None
    YArray = None
    YText = None
    y_apply_update = None
    y_encode_state_vector = None
    y_encode_state_as_update = None


def _b64e(b: bytes) -> str:
    return base64.b64encode(b or b"").decode("ascii")


def _b64d(s: str) -> bytes:
    try:
        return base64.b64decode((s or "").encode("ascii"))
    except Exception:
        return b""



# =========================
# Transport encryption (WS)
# =========================

try:
    from cryptography.hazmat.primitives.kdf.hkdf import HKDF
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.ciphers.aead import ChaCha20Poly1305
except Exception:
    HKDF = None
    hashes = None
    ChaCha20Poly1305 = None

_SW_AAD = b"sw-ws-v1"
_SW_HTTP_AAD = b"sw-http-v1"

def _sha256_hex(s: str) -> str:
    try:
        return hashlib.sha256((s or "").encode("utf-8")).hexdigest()
    except Exception:
        return ""

def _psk_from_activation_key(activation_key: str) -> bytes:
    if HKDF is None:
        raise RuntimeError("cryptography not installed")
    hkdf = HKDF(
        algorithm=hashes.SHA256(),
        length=32,
        salt=b"sw-psk-v1",
        info=b"sw-activation-psk",
    )
    return hkdf.derive((activation_key or "").encode("utf-8"))

def _session_key(psk: bytes, cnonce: bytes, snonce: bytes) -> bytes:
    hkdf = HKDF(
        algorithm=hashes.SHA256(),
        length=32,
        salt=cnonce + snonce,
        info=b"sw-session-v1",
    )
    return hkdf.derive(psk)

def _aead_encrypt(key: bytes, seq: int, plaintext: bytes, aad: bytes = _SW_AAD) -> bytes:
    if ChaCha20Poly1305 is None:
        raise RuntimeError("cryptography not installed")
    nonce = b"\x00\x00\x00\x00" + struct.pack(">Q", int(seq))
    return ChaCha20Poly1305(key).encrypt(nonce, plaintext, aad)

def _aead_decrypt(key: bytes, seq: int, ciphertext: bytes, aad: bytes = _SW_AAD) -> bytes:
    if ChaCha20Poly1305 is None:
        raise RuntimeError("cryptography not installed")
    nonce = b"\x00\x00\x00\x00" + struct.pack(">Q", int(seq))
    return ChaCha20Poly1305(key).decrypt(nonce, ciphertext, aad)

def _http_encrypt_payload(activation_key: str, payload: dict) -> dict:
    if ChaCha20Poly1305 is None:
        raise RuntimeError("cryptography not installed")
    key = _psk_from_activation_key(activation_key)
    nonce = os.urandom(12)
    pt = json.dumps(payload or {}, ensure_ascii=False).encode("utf-8")
    ct = ChaCha20Poly1305(key).encrypt(nonce, pt, _SW_HTTP_AAD)
    return {"sw_enc": {"nonce": _b64e(nonce), "ct": _b64e(ct)}}

def _http_decrypt_payload(activation_key: str, data: dict) -> dict:
    if ChaCha20Poly1305 is None:
        raise RuntimeError("cryptography not installed")
    if not isinstance(data, dict) or "sw_enc" not in data:
        return data
    enc = data.get("sw_enc") or {}
    nonce = _b64d(enc.get("nonce") or "")
    ct = _b64d(enc.get("ct") or "")
    key = _psk_from_activation_key(activation_key)
    pt = ChaCha20Poly1305(key).decrypt(nonce, ct, _SW_HTTP_AAD)
    return json.loads(pt.decode("utf-8"))

class CrdtRealtimeSync(QtCore.QObject):
    """Realtime collaboration via Y-CRDT (y-py) over server websocket.

    Server contract (current PlotWriter CRDT server):
      - WS endpoint: /ws?token=...&activation_key=...
      - Client sends:
            {"type":"crdt_subscribe","game_id":"...","state_vector_b64":"<b64>"}
      - Server replies:
            {"type":"crdt_subscribed","game_id":"...","role":"admin|master|reader","update_b64":"<b64>","state_vector_b64":"<b64>"}
        and later broadcasts:
            {"type":"crdt_update","game_id":"...","update_b64":"<b64>","user_id":"..."}
      - Client may send:
            {"type":"crdt_update","game_id":"...","update_b64":"<b64>"}
        server responds with:
            {"type":"crdt_ack","game_id":"...","state_vector_b64":"<b64>"}
    """

    statusChanged = QtCore.Signal(str)
    stateReceived = QtCore.Signal(dict, int)   # state, dummy revision
    remoteApplied = QtCore.Signal()
    errorOccurred = QtCore.Signal(str)

    def __init__(self, ctx: OnlineContext, parent=None):
        super().__init__(parent)
        self.ctx = ctx

        self.ws = QWebSocket()
        self.ws.connected.connect(self._on_connected)
        self.ws.disconnected.connect(self._on_disconnected)
        self.ws.textMessageReceived.connect(self._on_text)

        self._connected = False
        self._subscribed = False
        self._my_role = None
        self._last_server_sv = None

        self.ydoc = YDoc() if YDoc is not None else None

        # Last exported snapshot (server-like dict) to compute diffs.
        self._last_snapshot: Dict[str, Any] = {}

        # Debounce exporting local changes into YDoc
        self._dirty = False
        self._debounce = QtCore.QTimer(self)
        self._debounce.setInterval(200)
        self._debounce.setSingleShot(True)
        self._debounce.timeout.connect(self._flush_if_needed)

        # Prevent echo: when applying remote update, don't treat it as local user edit.
        self._applying_remote = False

        # Hook YDoc update observer: send incremental updates as they happen.
        # y-py callbacks differ between versions; support multiple APIs.
        self._y_sub = None
        self._has_update_observer = False
        self._last_sent_sv: Optional[bytes] = None  # state-vector at last local send

        if self.ydoc is not None:
            for meth in ("observe_update_v1", "observe_update"):
                fn = getattr(self.ydoc, meth, None)
                if not fn:
                    continue
                try:
                    self._y_sub = fn(self._on_ydoc_update)
                    self._has_update_observer = True
                    break
                except Exception:
                    self._y_sub = None
                    self._has_update_observer = False

        # Provided by MainWindow at runtime:
        #   self.get_current_state -> callable returning dict (project_to_server_state)
        self.get_current_state = None

        # Mirror for old code paths
        self.server_state: Dict[str, Any] = {}
        self.revision: int = 0  # not used in CRDT mode

        # Passive local-change detector: if UI code forgets to call mark_dirty(),
        # we still pick up edits and publish them. Cheap and reliable.
        self._last_local_hash = ""
        self._poll = QtCore.QTimer(self)
        self._poll.setInterval(400)
        self._poll.timeout.connect(self._poll_local_state)


        # ---- Transport encryption state (sw_hello / sw_msg) ----
        self._enc_ready = False
        self._psk = None
        self._sk = None
        self._cnonce = None
        self._send_seq = 0
        self._recv_seq = 0
        self._await_auth = False

    # ---- WS wiring ----
    def start(self):
        if self.ydoc is None:
            self.errorOccurred.emit("CRDT недоступен: установи пакет y-py (pip install y-py).")
            return

        base_ws = server_url_to_ws(self.ctx.server_url).rstrip("/")
        url = f"{base_ws}/ws"

        # Reset transport state before every connection attempt
        self._enc_ready = False
        self._psk = None
        self._sk = None
        self._cnonce = None
        self._send_seq = 0
        self._recv_seq = 0
        self._await_auth = False

        self.statusChanged.emit("CRDT: подключение…")
        self.ws.open(QNetworkRequest(QtCore.QUrl(url)))
        try:
            self._poll.start()
        except Exception:
            pass

    def stop(self):
        try:
            self.ws.close()
        except Exception:
            pass

    def _on_connected(self):
        self._connected = True
        self._subscribed = False
        self._my_role = None
        self._last_server_sv = None
        self.statusChanged.emit("CRDT: подключено, рукопожатие…")
        # Slight delay avoids rare cases where sending immediately after open() gets dropped.
        QtCore.QTimer.singleShot(50, self._start_handshake)

    def _on_disconnected(self):
        self._connected = False
        self._subscribed = False
        self._my_role = None
        self._last_server_sv = None
        self.statusChanged.emit("CRDT: отключено")

    def _send_plain(self, payload: dict):
        try:
            self.ws.sendTextMessage(json.dumps(payload, ensure_ascii=False))
        except Exception as e:
            self.errorOccurred.emit(f"CRDT send error: {e}")

    def _start_handshake(self):
        """Start encrypted transport handshake (sw_hello)."""
        if not self._connected:
            return
        act = getattr(self.ctx, 'activation_key', None) or ""
        if not act.strip():
            self.statusChanged.emit("CRDT: ошибка: activation_key не задан")
            self.errorOccurred.emit("CRDT: activation_key не задан (нужен для шифрования)")
            try:
                self.ws.close()
            except Exception:
                pass
            return
        if HKDF is None or ChaCha20Poly1305 is None:
            self.statusChanged.emit("CRDT: ошибка: нет cryptography")
            self.errorOccurred.emit("CRDT: модуль cryptography не установлен (нужен для шифрования)")
            try:
                self.ws.close()
            except Exception:
                pass
            return
        try:
            self._psk = _psk_from_activation_key(act)
            self._cnonce = os.urandom(32)
            key_id = _sha256_hex(act)
            mac = hmac.new(self._psk, b"hello" + self._cnonce, digestmod="sha256").digest()
            hello = {
                "type": "sw_hello",
                "key_id": key_id,
                "cnonce": _b64e(self._cnonce),
                "mac": _b64e(mac),
            }
            self._send_plain(hello)
        except Exception as e:
            self.statusChanged.emit(f"CRDT: ошибка рукопожатия: {e}")
            self.errorOccurred.emit(f"CRDT handshake error: {e}")
            try:
                self.ws.close()
            except Exception:
                pass

    def _send_payload(self, payload: dict):
        """Send payload over WS using encrypted sw_msg frames when ready."""
        if not self._connected:
            return
        if self._enc_ready and self._sk:
            try:
                pt = json.dumps(payload, ensure_ascii=False).encode('utf-8')
                ct = _aead_encrypt(self._sk, self._send_seq, pt, aad=_SW_AAD)
                wrap = {"type": "sw_msg", "seq": self._send_seq, "ct": _b64e(ct)}
                self.ws.sendTextMessage(json.dumps(wrap, ensure_ascii=False))
                self._send_seq += 1
            except Exception as e:
                self.errorOccurred.emit(f"CRDT encrypt/send error: {e}")
        else:
            # Before encryption is ready, only plaintext handshake frames are allowed.
            return

    def _handle_sw_hello_ok(self, msg: dict):
        try:
            if not self._psk or not self._cnonce:
                raise RuntimeError("handshake state missing")
            snonce = _b64d(msg.get('snonce') or '')
            mac2 = _b64d(msg.get('mac') or '')
            exp2 = hmac.new(self._psk, b"server" + self._cnonce + snonce, digestmod="sha256").digest()
            if not hmac.compare_digest(mac2, exp2):
                raise RuntimeError("bad server mac")
            self._sk = _session_key(self._psk, self._cnonce, snonce)
            self._enc_ready = True
            self._send_seq = 0
            self._recv_seq = 0
            self.statusChanged.emit("CRDT: шифрование активно, авторизация…")

            tok = getattr(self.ctx, 'access_token', None) or ""
            if not tok.strip():
                raise RuntimeError("access_token missing")
            self._await_auth = True
            self._send_payload({"type": "auth", "access_token": tok})
        except Exception as e:
            self.statusChanged.emit(f"CRDT: ошибка шифрования: {e}")
            self.errorOccurred.emit(f"CRDT: {e}")
            try:
                self.ws.close()
            except Exception:
                pass

    def _unwrap_if_encrypted(self, outer: dict) -> Optional[dict]:
        """If outer is sw_msg, decrypt and return inner payload dict."""
        if (outer or {}).get('type') != 'sw_msg':
            return outer
        if not self._enc_ready or not self._sk:
            return None
        try:
            seq = int(outer.get('seq', -1))
            if seq != int(self._recv_seq):
                raise RuntimeError(f"bad seq (got {seq}, expected {self._recv_seq})")
            ct = _b64d(outer.get('ct') or '')
            pt = _aead_decrypt(self._sk, self._recv_seq, ct, aad=_SW_AAD)
            self._recv_seq += 1
            inner = json.loads(pt.decode('utf-8'))
            return inner if isinstance(inner, dict) else None
        except Exception as e:
            self.errorOccurred.emit(f"CRDT decrypt error: {e}")
            return None

    def _send_subscribe(self):
        log_event("CRDT_SEND_SUBSCRIBE", {"game_id": getattr(self.ctx, "game_id", None), "sv_len": len(self._last_sent_sv or b"")})
        # Wait until encrypted auth is confirmed (auth_ok)
        if getattr(self, '_await_auth', False):
            return

        if not self._connected or self._subscribed or self.ydoc is None:
            return
        if not getattr(self.ctx, 'game_id', None):
            self.errorOccurred.emit('CRDT: game_id не передан (проверь параметры запуска --game_id=...)')
            self.statusChanged.emit('CRDT: ошибка (нет game_id)')
            return
        try:
            sv = y_encode_state_vector(self.ydoc)
            msg = {"type": "crdt_subscribe", "game_id": self.ctx.game_id, "state_vector_b64": _b64e(sv)}
            self._send_payload(msg)
        except Exception as e:
            self.errorOccurred.emit(f"CRDT subscribe error: {e}")


    def _on_text(self, text: str):
        log_event("CRDT_WS_RECV_RAW", {"n": len(text or "")})
        try:
            outer = json.loads(text or "{}")
        except Exception:
            return

        # Handshake (plaintext) frames
        if isinstance(outer, dict) and outer.get("type") == "sw_hello_ok":
            self._handle_sw_hello_ok(outer)
            return

        # Unwrap encrypted frames
        msg = self._unwrap_if_encrypted(outer if isinstance(outer, dict) else {})
        if not isinstance(msg, dict):
            return

        try:
            log_event("CRDT_WS_RECV", {"t": msg.get("type"), "keys": list(msg.keys())})
        except Exception:
            pass

        t = msg.get("type")

        # Server-side error
        if t == 'error':
            detail = msg.get('detail') or msg.get('message') or str(msg)
            self._subscribed = False
            self.statusChanged.emit(f"CRDT: ошибка: {detail}")
            self.errorOccurred.emit(f"CRDT: {detail}")
            return

        # Auth ok: after encrypted handshake we must auth before subscribe
        if t == 'auth_ok':
            if self._await_auth:
                self._await_auth = False
                QtCore.QTimer.singleShot(0, self._send_subscribe)
            return

        # Initial sync after subscribe
        if t == "crdt_subscribed":
            upd_b64 = msg.get("update_b64") or msg.get("update") or ""
            upd = _b64d(upd_b64)
            if upd:
                log_event("CRDT_APPLY_REMOTE", {"initial": True, "upd_len": len(upd)})
                self._apply_remote_update(upd, initial=True)

            # Server tells our role; used to gate local edits
            self._my_role = msg.get("role") or getattr(self, "_my_role", None)

            # We are subscribed even if update was empty
            self._subscribed = True
            self.statusChanged.emit("CRDT: подключено")
            try:
                if not self._poll.isActive():
                    self._poll.start()
            except Exception:
                pass
            return

        # Legacy initial sync message name (older builds)
        if t == 'crdt_sync':
            upd_b64 = msg.get('update_b64') or msg.get('update') or ''
            upd = _b64d(upd_b64)
            if upd:
                log_event("CRDT_APPLY_REMOTE", {"initial": True, "upd_len": len(upd)})
                self._apply_remote_update(upd, initial=True)
            self._subscribed = True
            self.statusChanged.emit('CRDT: подключено')
            return

        # Incremental update from server
        if t == "crdt_update":
            upd_b64 = msg.get("update_b64") or msg.get("update") or msg.get("data") or ""
            upd = _b64d(upd_b64)
            if upd:
                log_event("CRDT_APPLY_REMOTE", {"initial": False, "upd_len": len(upd)})
                self._apply_remote_update(upd, initial=False)
            return

        # Ack after our update (optional)
        if t == "crdt_ack":
            sv_b64 = msg.get("state_vector_b64") or ""
            if sv_b64:
                self._last_server_sv = _b64d(sv_b64)
            return

        # Ignore other server chatter
        if t in ("hello", "subscribed", "pong"):
            return

    def _send_update_bytes(self, upd: bytes):
        if self._applying_remote:
            return
        if not upd:
            return
        log_event("CRDT_WS_SEND", {"type": "crdt_update", "upd_len": len(upd)})
        if not self._connected or not self._subscribed:
            return
        # If role is reader, do not push edits (server will reject anyway)
        role = getattr(self, "_my_role", None)
        if role not in (None, "admin", "master"):
            return
        try:
            msg = {"type": "crdt_update", "game_id": self.ctx.game_id, "update_b64": _b64e(upd)}
            self._send_payload(msg)
        except Exception as e:
            self.errorOccurred.emit(f"CRDT send error: {e}")

    def _on_ydoc_update(self, evt):
        # y-py version differences:
        #  - some pass raw bytes
        #  - some pass an event object with .update (bytes)
        #  - some pass a tuple-like container
        if self._applying_remote:
            return
        upd = None

        try:
            if isinstance(evt, (bytes, bytearray)):
                upd = bytes(evt)
            elif hasattr(evt, "update"):
                u = getattr(evt, "update")
                if isinstance(u, (bytes, bytearray)):
                    upd = bytes(u)
            elif isinstance(evt, tuple) and evt:
                # y-py 0.6.x passes (txn, update) to observe_update_v1
                # Some builds may pass just (update,) or other tuple-like containers.
                if len(evt) >= 2 and isinstance(evt[1], (bytes, bytearray)):
                    upd = bytes(evt[1])
                elif isinstance(evt[0], (bytes, bytearray)):
                    upd = bytes(evt[0])
        except Exception:
            upd = None

        if not upd:
            # No update extracted. Don't spam; just silently ignore.
            return

        self._send_update_bytes(upd)

    # ---- Public API used by MainWindow ----

    def mark_dirty(self):
        self._dirty = True
        self._debounce.start()



    def _poll_local_state(self):
        """Detect local project changes even if UI forgot to call mark_dirty()."""
        if not self._subscribed:
            return
        if self._applying_remote:
            return
        if not callable(self.get_current_state):
            return
        try:
            snap = self.get_current_state() or {}
            # stable hash of snapshot
            s = json.dumps(snap, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
            h = hashlib.sha1(s.encode("utf-8")).hexdigest()
            if h != getattr(self, "_last_local_hash", ""):
                self._last_local_hash = h
                log_event("CRDT_POLL_DIRTY", {"changed": True})
                self.mark_dirty()
        except Exception as e:
            log_event("CRDT_POLL_ERR", {"err": str(e)})
    def flush_full(self, project: "Project"):
        if self.ydoc is None:
            return
        self._dirty = False
        snap = project_to_server_state(project)
        self._apply_snapshot_to_ydoc(snap, force=True)

    # ---- Internals ----
    def _flush_if_needed(self):
        if not self._dirty:
            return
        self._dirty = False
        if not callable(self.get_current_state) or self.ydoc is None:
            return
        try:
            snap = self.get_current_state() or {}
            self._apply_snapshot_to_ydoc(snap, force=False)
        except Exception as e:
            self.errorOccurred.emit(f"CRDT flush error: {e}")

    def _apply_remote_update(self, update: bytes, initial: bool):
        if self.ydoc is None or y_apply_update is None:
            return
        self._applying_remote = True
        try:
            y_apply_update(self.ydoc, update)

            state = self._export_ydoc_to_server_state()
            log_event("CRDT_EXPORT_STATE", {"keys": list(state.keys()) if isinstance(state, dict) else None})
            self.server_state = state
            if initial:
                self._subscribed = True
                self.statusChanged.emit("CRDT: подключено")
                self.stateReceived.emit(state, 0)
            else:
                self.remoteApplied.emit()

            self._last_snapshot = state
        
            try:
                s = json.dumps(state, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                self._last_local_hash = hashlib.sha1(s.encode("utf-8")).hexdigest()
            except Exception:
                pass
        except Exception as e:
            self.errorOccurred.emit(f"CRDT apply error: {e}")
        finally:
            self._applying_remote = False

    def _apply_snapshot_to_ydoc(self, snap: Dict[str, Any], force: bool):
        if self.ydoc is None:
            return
        # translate existing diff ops to CRDT writes
        ops = diff_state_to_ops({} if force else (self._last_snapshot or {}), snap)
        if not ops:
            return
        with self.ydoc.begin_transaction() as txn:
            root = self._get_or_create_root(txn)
            self._apply_ops_as_crdt(txn, root, ops, snap)

        # Update our local baseline so we don't re-send the same ops forever.
        self._last_snapshot = snap
        self.server_state = snap

        # Push local changes to server explicitly (do not rely on y-py observers).
        try:
            if y_encode_state_as_update is not None and y_encode_state_vector is not None and self.ydoc is not None:
                prev_sv = getattr(self, "_last_sent_sv", None)
                upd = y_encode_state_as_update(self.ydoc, prev_sv) if prev_sv else y_encode_state_as_update(self.ydoc)
                self._last_sent_sv = y_encode_state_vector(self.ydoc)
                if upd:
                    log_event("CRDT_SEND_UPDATE", {"upd_len": len(upd)})
                    self._send_update_bytes(upd)
        except Exception as e:
            log_event("CRDT_SEND_UPDATE_ERR", {"err": str(e)})

        self._last_snapshot = snap

    def _get_or_create_root(self, txn):
        root = self.ydoc.get_map("root")
        # Ensure collections exist
        for k in ("characters", "factions", "locations", "stories", "hooks", "goals", "custom_titles"):
            v = root.get(k)
            if v is None:
                root.set(txn, k, YMap({}))
        if root.get("schema_version") is None:
            root.set(txn, "schema_version", int(SCHEMA_VERSION))
        return root

    def _apply_ops_as_crdt(self, txn, root, ops: List[Dict[str, Any]], snap: Dict[str, Any]):
        for op in ops:
            o = op.get("op")
            if o == "set" and op.get("path") == "schema_version":
                root.set(txn, "schema_version", int(op.get("value", SCHEMA_VERSION)))
                continue

            if o == "set" and op.get("path") == "custom_titles":
                ct = op.get("value") or {}
                m = root.get("custom_titles")
                if not isinstance(m, YMap):
                    m = YMap({})
                    root.set(txn, "custom_titles", m)
                m.set(txn, "characters", json.dumps(list(ct.get("characters") or []), ensure_ascii=False))
                m.set(txn, "factions", json.dumps(list(ct.get("factions") or []), ensure_ascii=False))
                continue

            if o == "delete":
                path = op.get("path", "")
                if "." not in path:
                    continue
                col, eid = path.split(".", 1)
                cm = root.get(col)
                if not isinstance(cm, YMap):
                    continue
                try:
                    cm.pop(txn, eid)
                except Exception:
                    cm.set(txn, eid, None)
                continue

            if o == "upsert_entity":
                col = op.get("collection") or ""
                eid = op.get("id") or ""
                if not col or not eid:
                    continue
                ent = (snap.get(col) or {}).get(eid) or op.get("value") or {}
                if not isinstance(ent, dict):
                    continue

                cm = root.get(col)
                if not isinstance(cm, YMap):
                    cm = YMap({})
                    root.set(txn, col, cm)

                em = cm.get(eid)
                if not isinstance(em, YMap):
                    em = YMap({})
                    cm.set(txn, eid, em)

                for k, v in ent.items():
                    if v is None:
                        continue
                    if k in ("story_public", "story_private", "description", "label", "title", "name") and isinstance(v, str):
                        yt = em.get(k)
                        if not isinstance(yt, YText):
                            yt = YText("")
                            em.set(txn, k, yt)
                        try:
                            cur = yt.to_string()
                            if cur != v:
                                yt.delete(txn, 0, len(cur))
                                if v:
                                    yt.insert(txn, 0, v)
                        except Exception:
                            em.set(txn, k, v)
                    elif isinstance(v, (str, int, float, bool)):
                        em.set(txn, k, v)
                    else:
                        # store complex values as json string for stability
                        em.set(txn, k, json.dumps(v, ensure_ascii=False))

    def _export_ydoc_to_server_state(self) -> Dict[str, Any]:
        if self.ydoc is None:
            return {}

        out: Dict[str, Any] = {"schema_version": SCHEMA_VERSION}
        root = self.ydoc.get_map("root")
        try:
            sv = root.get("schema_version")
            if isinstance(sv, int):
                out["schema_version"] = sv
        except Exception:
            pass

        def read_coll(col: str) -> Dict[str, Any]:
            cm = root.get(col)
            if not isinstance(cm, YMap):
                return {}
            res: Dict[str, Any] = {}
            try:
                items = cm.items()
            except Exception:
                items = []
            for eid, em in items:
                if not isinstance(eid, str) or not isinstance(em, YMap):
                    continue
                ent: Dict[str, Any] = {}
                try:
                    kvs = em.items()
                except Exception:
                    kvs = []
                for k, v in kvs:
                    if isinstance(v, YText):
                        try:
                            ent[k] = v.to_string()
                        except Exception:
                            ent[k] = ""
                    elif isinstance(v, str):
                        # try json decode
                        if (v.startswith("{") and v.endswith("}")) or (v.startswith("[") and v.endswith("]")):
                            try:
                                ent[k] = json.loads(v)
                                continue
                            except Exception:
                                pass
                        ent[k] = v
                    else:
                        ent[k] = v
                res[eid] = ent
            return res

        for col in ("characters", "factions", "locations", "stories", "hooks", "goals"):
            out[col] = read_coll(col)

        # custom titles
        try:
            ct = root.get("custom_titles")
            if isinstance(ct, YMap):
                chars = ct.get("characters")
                facs = ct.get("factions")
                if isinstance(chars, str):
                    try:
                        chars = json.loads(chars)
                    except Exception:
                        chars = []
                if isinstance(facs, str):
                    try:
                        facs = json.loads(facs)
                    except Exception:
                        facs = []
                out["custom_titles"] = {"characters": chars or [], "factions": facs or []}
        except Exception:
            pass

        return out


# =========================
# Main Window
# =========================

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self, online: Optional[OnlineContext] = None):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} {get_app_version()} (schema v{SCHEMA_VERSION})")
        self.resize(1500, 880)

        self.project = Project()
        self.online = online or OnlineContext()
        self._dirty = False
        self._applying_remote = False
        self.sync_api: Optional[SyncApiClient] = None
        self.sync: Optional[RealtimeSync] = None
        self._sync_status = ""
        # Keep custom field titles consistent within this project
        self.project.sync_custom_field_titles_from_objects()
        self._auto_timer = QtCore.QTimer(self)
        self._auto_timer.setInterval(60_000)
        self._auto_timer.timeout.connect(self.auto_version_tick)
        self._auto_timer.start()

        self.build_menus()

        self.tabs = QtWidgets.QTabWidget()
        self.setCentralWidget(self.tabs)

        self.scene = FieldScene(self.project)
        self.scene.on_open_character = self.open_character
        self.scene.on_open_hooks = self.open_hooks_for_char
        self.scene.on_open_goals = self.open_goals_for_char
        self.scene.on_open_hook_by_id = self.open_hook_by_id

        # background context menu callbacks
        self.scene.on_create_character_at = self.create_character_at
        self.scene.on_create_faction_at = self.create_faction_at


        self.scene.on_open_faction = self.open_faction
        self.scene.on_delete_character = self.delete_character
        self.scene.on_delete_faction = self.delete_faction
        self.scene.on_project_changed = self.on_local_change
        self.field_view = FieldView(self.scene)
        field_wrap = QtWidgets.QWidget()
        fw = QtWidgets.QVBoxLayout(field_wrap)
        top = QtWidgets.QHBoxLayout()
        self.btn_relayout = QtWidgets.QPushButton("Переразложить")
        top.addWidget(self.btn_relayout)
        top.addStretch(1)
        fw.addLayout(top)
        fw.addWidget(self.field_view, 1)

        # Zoom bar (bottom)
        zoom_row = QtWidgets.QHBoxLayout()
        self.lbl_zoom = QtWidgets.QLabel("Масштаб: 100%")
        self.lbl_zoom.setObjectName("Tip")
        self.zoom_slider = QtWidgets.QSlider(QtCore.Qt.Horizontal)
        self.zoom_slider.setRange(25, 300)
        self.zoom_slider.setValue(100)
        self.zoom_slider.setSingleStep(5)
        self.zoom_slider.setPageStep(25)
        self.zoom_slider.setFixedHeight(22)
        self.zoom_slider.setToolTip("Масштаб поля")
        zoom_row.addWidget(self.lbl_zoom)
        zoom_row.addWidget(self.zoom_slider, 1)
        fw.addLayout(zoom_row)

        self._zoom_sync_lock = False

        def on_slider(val: int):
            if self._zoom_sync_lock:
                return
            self._zoom_sync_lock = True
            try:
                self.field_view.set_zoom_percent(val)
                self.lbl_zoom.setText(f"Масштаб: {self.field_view.current_zoom_percent()}%")
            finally:
                self._zoom_sync_lock = False

        def on_zoom_changed(val: int):
            if self._zoom_sync_lock:
                return
            self._zoom_sync_lock = True
            try:
                self.zoom_slider.setValue(val)
                self.lbl_zoom.setText(f"Масштаб: {val}%")
            finally:
                self._zoom_sync_lock = False

        self.zoom_slider.valueChanged.connect(on_slider)
        self.field_view.zoomChanged.connect(on_zoom_changed)
        self.btn_relayout.clicked.connect(self.relayout)

        self.characters_tab = CharactersTab(self.project, self.open_character, change_cb=self.on_local_change)
        self.factions_tab = SimpleListTab("Фракция", self.project, "faction",
                                          open_cb=self.open_faction,
                                          create_cb=self.create_faction)
        self.locations_tab = SimpleListTab("Локация", self.project, "location",
                                           open_cb=self.open_location,
                                           create_cb=self.create_location)
        self.stories_tab = SimpleListTab("Сюжет", self.project, "story",
                                         open_cb=self.open_story,
                                         create_cb=self.create_story)

        self.tabs.addTab(field_wrap, "Поле")
        self.tabs.addTab(self.characters_tab, "Персонажи")
        self.tabs.addTab(self.factions_tab, "Фракции")
        self.tabs.addTab(self.locations_tab, "Локации")
        self.tabs.addTab(self.stories_tab, "Сюжеты")

        self.scene.rebuild()

        self.statusBar().showMessage("Готово.")
        if self.online and self.online.mode == "online" and self.online.server_url and self.online.game_id:
            self.init_online_sync()


    # ===== Online sync =====
    def init_online_sync(self):
        # Build API + realtime sync (CRDT preferred)
        try:
            self.sync_api = SyncApiClient(self.online)

            # Prefer CRDT if y-py is available; otherwise fall back to legacy op-sync.
            if YDoc is not None:
                self.sync = CrdtRealtimeSync(self.online, self)
            else:
                self.sync = RealtimeSync(self.sync_api, self.online, self)

            # UI callbacks
            self.sync.statusChanged.connect(self._on_sync_status)
            self.sync.stateReceived.connect(self._on_sync_state)
            self.sync.remoteApplied.connect(self._on_sync_remote_applied)
            self.sync.errorOccurred.connect(self._on_sync_error)

            # provide current state callback
            self.sync.get_current_state = lambda: project_to_server_state(self.project)

            self.sync.start()
        except Exception as e:
            log_event("SYNC_INIT_ERROR", err=str(e))
            self.statusBar().showMessage(f"Sync init error: {e}")
            self.sync = None

    def _on_sync_status(self, s: str):
        self._sync_status = s
        self.statusBar().showMessage(s)

    def _on_sync_error(self, s: str):
        # No popups for transient network stuff. Keep it readable.
        self.statusBar().showMessage(f"Sync: {s}")

    def _on_sync_state(self, state: dict, revision: int):
        # Initial load from server
        self._applying_remote = True
        try:
            self.apply_server_state(state, quiet=True)
            # After loading, keep baseline for diff
            if self.sync:
                self.sync.server_state = state or {}
                self.sync.revision = int(revision)
            self._dirty = False
        finally:
            self._applying_remote = False

    def _on_sync_remote_applied(self):
        # Remote event already applied to sync.server_state, update project quietly
        if not self.sync:
            return
        self._applying_remote = True
        try:
            self.apply_server_state(self.sync.server_state, quiet=True)
        finally:
            self._applying_remote = False

    def apply_server_state(self, state: dict, quiet: bool = False):
        # Convert and load into Project
        cur = server_state_to_editor_current(state or {})
        # apply custom titles if present
        titles = (state or {}).get("custom_titles") or {}
        if isinstance(titles, dict):
            ch_t = titles.get("characters")
            fa_t = titles.get("factions")
            if isinstance(ch_t, list) and len(ch_t) == 5:
                self.project.character_custom_titles = [norm_spaces(x) for x in ch_t]
            if isinstance(fa_t, list) and len(fa_t) == 5:
                self.project.faction_custom_titles = [norm_spaces(x) for x in fa_t]

        self.project.load_state(cur)
        # Enforce global titles (important for multi-client consistency)
        self.project.apply_custom_field_titles()
        self.project.rebuild_dictionaries()


        if quiet:
            self.refresh_all_quiet()
        else:
            self.refresh_all()

    def on_local_change(self):
        if self._applying_remote:
            return
        self._dirty = True
        if self.sync and self.online.mode == "online":
            self.sync.mark_dirty()

    def refresh_all_quiet(self):
        # Try hard not to make the UI blink when remote updates arrive
        w = self.centralWidget()
        if w:
            w.setUpdatesEnabled(False)
        self.setUpdatesEnabled(False)
        try:
            self.refresh_all()
        finally:
            if w:
                w.setUpdatesEnabled(True)
            self.setUpdatesEnabled(True)
            self.update()

    # ===== Save behavior =====
    def _push_online_if_needed(self, comment: str = ""):
        if not (self.sync and self.sync_api and self.online.mode == "online"):
            return
        try:
            # Make sure positions are up to date before syncing
            self.scene.sync_visuals_to_project()
            self.sync.flush_full(self.project)
            # optional server-side version snapshot
            try:
                self.sync_api.save_version(comment=comment or "")
            except Exception:
                pass
        except Exception as e:
            self.statusBar().showMessage(f"Sync save failed: {e}")

    def closeEvent(self, event: QtGui.QCloseEvent):
        # Save on close (local + server if online). No drama, no prompts.
        try:
            if not self.project.file_path:
                # autosave near executable / script
                base = os.path.dirname(os.path.abspath(__file__))
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                self.project.file_path = os.path.join(base, f"autosave_{ts}.swproj.json")
            # always sync visuals
            self.scene.sync_visuals_to_project()
            # local save
            self.project.save_to_path(self.project.file_path, kind="autosave", author=self.online.username or "local", comment="close")
            # online save
            self._push_online_if_needed(comment="close")
        except Exception:
            pass
        event.accept()

    # ===== Menus =====
    def build_menus(self):
        mb = self.menuBar()

        # File menu exactly as requested (and only those items)
        m_file = mb.addMenu("Файл")
        act_new = m_file.addAction("Новый проект")
        act_open = m_file.addAction("Открыть…")
        act_save = m_file.addAction("Сохранить")
        act_saveas = m_file.addAction("Сохранить как…")
        m_file.addSeparator()
        act_history = m_file.addAction("История версий…")
        m_file.addSeparator()
        act_exit = m_file.addAction("Выход")

        act_new.triggered.connect(self.new_project)
        act_open.triggered.connect(self.open_project)
        act_save.triggered.connect(self.save_project)
        act_saveas.triggered.connect(self.save_project_as)
        act_history.triggered.connect(self.open_history)
        act_exit.triggered.connect(self.close)

        # Export menu next to File
        m_exp = mb.addMenu("Экспорт персонажей")
        act_export_sel = m_exp.addAction("Экспорт выбранных персонажей (PDF)…")
        act_export_all = m_exp.addAction("Экспорт всех персонажей (PDF)…")
        act_export_sel_docx = m_exp.addAction("Экспорт выбранных персонажей (DOCX)…")
        act_export_all_docx = m_exp.addAction("Экспорт всех персонажей (DOCX)…")

        act_export_settings = m_exp.addAction("????????? ?????????")
        act_export_sel.triggered.connect(self.export_selected_characters_pdf)
        act_export_all.triggered.connect(self.export_all_characters_pdf)
        act_export_sel_docx.triggered.connect(self.export_selected_characters_docx)
        act_export_all_docx.triggered.connect(self.export_all_characters_docx)
        act_export_settings.triggered.connect(self.open_export_settings)


    def open_export_settings(self):
        dlg = ExportSettingsDialog(self.project, change_cb=self.on_local_change, parent=self)
        dlg.exec()

    # ===== Navigation by hyperlink =====
    def open_by_name(self, name: str):
        obj = self.project.get_object_by_name(name)
        if not obj:
            return
        t, oid = obj
        if t == "character":
            self.open_character(oid)
        elif t == "faction":
            self.open_faction(oid)
        elif t == "location":
            self.open_location(oid)
        elif t == "story":
            self.open_story(oid)

    # ===== Create items from field context menu =====
    def create_character_at(self, pos: QtCore.QPointF):
        c = Character(name="Новый персонаж")
        c.visual.x, c.visual.y = float(pos.x()), float(pos.y())
        self.project.characters[c.id] = c
        self.open_character(c.id)  # keeps existing behavior (edit right away)
        self.project.rebuild_dictionaries()
        self.on_local_change()
        self.refresh_all()

    def create_faction_at(self, pos: QtCore.QPointF):
        f = Faction(name="Новая фракция")
        f.visual.x, f.visual.y = float(pos.x()), float(pos.y())
        self.project.factions[f.id] = f
        # Apply global custom field titles so the dialog shows consistent names
        self.project.apply_custom_field_titles()
        self.open_faction(f.id)
        self.project.rebuild_dictionaries()
        self.on_local_change()
        self.refresh_all()

    # ===== Open dialogs =====
    def open_character(self, char_id: str):
        c = self.project.characters.get(char_id)
        if not c or c.meta.is_deleted:
            return
        dlg = CharacterDialog(
            self.project,
            c,
            open_by_name_cb=self.open_by_name,
            open_hooks_cb=self.open_hooks_for_char,
            open_goals_cb=self.open_goals_for_char,
            change_cb=self._on_live_character_changed,
            parent=self
        )
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.project.rebuild_dictionaries()
            self.on_local_change()
            self.refresh_all()

    def open_goals_for_char(self, char_id: str):
        c = self.project.characters.get(char_id)
        if not c or c.meta.is_deleted:
            return
        dlg = GoalsDialog(self.project, focus_char_id=char_id, open_by_name_cb=self.open_by_name, parent=self)
        dlg.exec()
        self.on_local_change()
        self.refresh_all()


    def delete_character(self, char_id: str):
        c = self.project.characters.get(char_id)
        if not c or c.meta.is_deleted:
            return
        r = QtWidgets.QMessageBox.question(self, "Удаление", f"Удалить персонажа “{c.name}”? (мягкое удаление)")
        if r != QtWidgets.QMessageBox.Yes:
            return
        c.meta.is_deleted = True
        c.meta.touch("local")
        self.project.rebuild_dictionaries()
        self.on_local_change()
        self.refresh_all()

    def delete_faction(self, faction_id: str):
        f = self.project.factions.get(faction_id)
        if not f or f.meta.is_deleted:
            return
        r = QtWidgets.QMessageBox.question(self, "Удаление", f"Удалить фракцию “{f.name}”? (мягкое удаление)")
        if r != QtWidgets.QMessageBox.Yes:
            return
        f.meta.is_deleted = True
        f.meta.touch("local")
        self.project.rebuild_dictionaries()
        self.on_local_change()
        self.refresh_all()


    def open_faction(self, faction_id: str):
        f = self.project.factions.get(faction_id)
        if not f or f.meta.is_deleted:
            return
        dlg = FactionDialog(self.project, f, open_by_name_cb=self.open_by_name, change_cb=self._on_live_faction_changed, parent=self)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.project.rebuild_dictionaries()
            self.on_local_change()
            self.refresh_all()

    def _on_live_character_changed(self, char_id: str):
        # Update label without full rebuild to avoid flicker during live edits
        try:
            c = self.project.characters.get(char_id)
            it = self.scene.char_items.get(char_id) if hasattr(self, 'scene') else None
            if c and it and hasattr(it, 'text'):
                it.text.setText(c.name)
        except Exception:
            pass
        self.on_local_change()

    def _on_live_faction_changed(self, faction_id: str):
        try:
            f = self.project.factions.get(faction_id)
            it = self.scene.faction_items.get(faction_id) if hasattr(self, 'scene') else None
            if f and it and hasattr(it, 'text'):
                it.text.setText(f.name)
        except Exception:
            pass
        self.on_local_change()

    def open_location(self, location_id: str):
        l = self.project.locations.get(location_id)
        if not l or l.meta.is_deleted:
            return
        dlg = LocationDialog(self.project, l, open_by_name_cb=self.open_by_name, parent=self)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.project.rebuild_dictionaries()
            self.on_local_change()
            self.refresh_all()

    def open_story(self, story_id: str):
        s = self.project.stories.get(story_id)
        if not s or s.meta.is_deleted:
            return
        dlg = StoryDialog(self.project, s, open_by_name_cb=self.open_by_name, parent=self)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.project.rebuild_dictionaries()
            self.on_local_change()
            self.refresh_all()

    def open_hooks_for_char(self, char_id: str):
        dlg = HooksDialog(self.project, focus_char_id=char_id, parent=self)
        dlg.exec()
        self.refresh_all()

    def open_hook_by_id(self, hook_id: str):
        dlg = HooksDialog(self.project, focus_hook_id=hook_id, parent=self)
        dlg.exec()
        self.on_local_change()
        self.refresh_all()

    # ===== Create from tabs (unchanged) =====
    def create_faction(self):
        f = Faction(name="Новая фракция")
        f.visual.x, f.visual.y = 0.0, 0.0
        self.project.factions[f.id] = f
        self.on_local_change()
        self.open_faction(f.id)

    def create_location(self):
        l = Location(name="Новая локация")
        self.project.locations[l.id] = l
        self.on_local_change()
        self.open_location(l.id)

    def create_story(self):
        s = Story(name="Новый сюжет")
        self.project.stories[s.id] = s
        self.on_local_change()
        self.open_story(s.id)

    # ===== Refresh / layout =====
    def refresh_all(self):
        self.scene.sync_visuals_to_project()
        self.scene.rebuild()
        self.characters_tab.refresh_filters()
        self.characters_tab.refresh()
        self.factions_tab.refresh()
        self.locations_tab.refresh()
        self.stories_tab.refresh()

    def relayout(self):
        self.scene.auto_layout()
        self.on_local_change()
        self.refresh_all()

    # ===== Project operations =====
    def new_project(self):
        r = QtWidgets.QMessageBox.question(self, "Новый проект", "Создать новый проект? Несохранённые изменения пропадут.")
        if r != QtWidgets.QMessageBox.Yes:
            return
        self.project = Project()
        self.scene.project = self.project
        self.characters_tab.project = self.project
        self.factions_tab.project = self.project
        self.locations_tab.project = self.project
        self.stories_tab.project = self.project
        self.refresh_all()

    def open_project(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Открыть проект", "", "StoryWeaver Project (*.swproj *.json);;All files (*.*)")
        if not path:
            return
        try:
            self.project.open_from_path(path)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл:\n{e}")
            return
        self.scene.project = self.project
        self.characters_tab.project = self.project
        self.factions_tab.project = self.project
        self.locations_tab.project = self.project
        self.stories_tab.project = self.project
        self.refresh_all()

    def save_project(self):
        if not self.project.file_path:
            self.save_project_as()
            return
        try:
            self.scene.sync_visuals_to_project()
            author = self.online.username if (self.online and self.online.username) else "local"
            self.project.save_to_path(self.project.file_path, kind="manual", author=author, comment="")
            self._push_online_if_needed(comment="manual")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить:\n{e}")
            return
        self.statusBar().showMessage("Сохранено.", 3000)

    def save_project_as(self):
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Сохранить проект как", "", "StoryWeaver Project (*.swproj);;JSON (*.json)")
        if not path:
            return
        if not (path.endswith(".swproj") or path.endswith(".json")):
            path += ".swproj"
        try:
            self.scene.sync_visuals_to_project()
            author = self.online.username if (self.online and self.online.username) else "local"
            self.project.save_to_path(path, kind="manual", author=author, comment="")
            self._push_online_if_needed(comment="manual")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить:\n{e}")
            return
        self.statusBar().showMessage("Сохранено.", 3000)

    def open_history(self):
        dlg = HistoryDialog(self.project, parent=self)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.refresh_all()

    def auto_version_tick(self):
        if not self.project.file_path:
            self.project.versions.append(VersionEntry(kind="auto", author="local", comment="auto", state=self.project.snapshot_state()))
            return
        try:
            self.project.save_to_path(self.project.file_path, kind="auto", author="local", comment="auto")
            self.statusBar().showMessage("Автосохранение (версия создана).", 2000)
        except Exception:
            pass

    # ===== Export PDF =====
    def _default_title_map_for_mass(self) -> List[str]:
        titles = [""] * 5
        for c in self.project.alive_characters():
            for i in range(5):
                t = norm_spaces(c.custom_tag_fields[i].field_title)
                if t and not titles[i]:
                    titles[i] = t
            if any(titles):
                break
        return titles

    def export_selected_characters_pdf(self):
        cids = self.characters_tab.get_selected_character_ids()
        if not cids:
            QtWidgets.QMessageBox.information(self, "Экспорт", "Вкладка 'Персонажи': выдели одну или несколько строк для экспорта.")
            return
        chars = []
        for cid in cids:
            c = self.project.characters.get(cid)
            if c and not c.meta.is_deleted:
                chars.append(c)
        if not chars:
            return

        title_map = self._default_title_map_for_mass()
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Куда экспортировать PDF?")
        if not folder:
            return

        count = 0
        for c in chars:
            fname = safe_filename(c.name) + ".pdf"
            out_path = os.path.join(folder, fname)
            pdf_export_character(self.project, c, fields, out_path, font_name=self.project.export_font)
            count += 1

        QtWidgets.QMessageBox.information(self, "Готово", f"Экспортировано файлов: {count}")

    def export_all_characters_pdf(self):
        chars = self.project.alive_characters()
        if not chars:
            QtWidgets.QMessageBox.information(self, "Экспорт", "В проекте нет персонажей для экспорта.")
            return

        title_map = self._default_title_map_for_mass()
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Куда экспортировать PDF?")
        if not folder:
            return

        count = 0
        for c in chars:
            fname = safe_filename(c.name) + ".pdf"
            out_path = os.path.join(folder, fname)
            pdf_export_character(self.project, c, fields, out_path, font_name=self.project.export_font)
            count += 1

        QtWidgets.QMessageBox.information(self, "Готово", f"Экспортировано файлов: {count}")

    def export_selected_characters_docx(self):
        if Document is None:
            QtWidgets.QMessageBox.critical(self, "Ошибка", "Модуль python-docx не установлен.")
            return
        cids = self.characters_tab.get_selected_character_ids()
        if not cids:
            QtWidgets.QMessageBox.information(self, "Экспорт", "Вкладка 'Персонажи': выдели одну или несколько строк для экспорта.")
            return
        chars = []
        for cid in cids:
            c = self.project.characters.get(cid)
            if c and not c.meta.is_deleted:
                chars.append(c)
        if not chars:
            return

        title_map = self._default_title_map_for_mass()
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Куда экспортировать DOCX?")
        if not folder:
            return

        count = 0
        for c in chars:
            fname = safe_filename(c.name) + ".docx"
            out_path = os.path.join(folder, fname)
            docx_export_character(self.project, c, fields, out_path, font_name=self.project.export_font)
            count += 1

        QtWidgets.QMessageBox.information(self, "Готово", f"Экспортировано файлов: {count}")

    def export_all_characters_docx(self):
        if Document is None:
            QtWidgets.QMessageBox.critical(self, "Ошибка", "Модуль python-docx не установлен.")
            return
        chars = self.project.alive_characters()
        if not chars:
            QtWidgets.QMessageBox.information(self, "Экспорт", "В проекте нет персонажей для экспорта.")
            return

        title_map = self._default_title_map_for_mass()
        dlg = ExportFieldsDialog(title_map, defaults=self.project.export_defaults, parent=self)
        if dlg.exec() != QtWidgets.QDialog.Accepted:
            return
        fields = dlg.get_fields()

        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Куда экспортировать DOCX?")
        if not folder:
            return

        count = 0
        for c in chars:
            fname = safe_filename(c.name) + ".docx"
            out_path = os.path.join(folder, fname)
            docx_export_character(self.project, c, fields, out_path, font_name=self.project.export_font)
            count += 1

        QtWidgets.QMessageBox.information(self, "Готово", f"Экспортировано файлов: {count}")



# =========================
# Theme (match launcher)
# =========================

def apply_dark_theme(app: QtWidgets.QApplication):
    """
    Force a consistent dark theme on every machine (Windows theme can be... creative).
    Purely visual: no functional changes.
    """
    app.setStyle("Fusion")
    app.setFont(QtGui.QFont("Segoe UI", 10))

    pal = QtGui.QPalette()
    pal.setColor(QtGui.QPalette.Window, QtGui.QColor("#0f1115"))
    pal.setColor(QtGui.QPalette.WindowText, QtGui.QColor("#e6e6e6"))
    pal.setColor(QtGui.QPalette.Base, QtGui.QColor("#161a22"))
    pal.setColor(QtGui.QPalette.AlternateBase, QtGui.QColor("#121622"))
    pal.setColor(QtGui.QPalette.Text, QtGui.QColor("#e6e6e6"))
    pal.setColor(QtGui.QPalette.Button, QtGui.QColor("#161a22"))
    pal.setColor(QtGui.QPalette.ButtonText, QtGui.QColor("#e6e6e6"))
    pal.setColor(QtGui.QPalette.Highlight, QtGui.QColor("#2b64ff"))
    pal.setColor(QtGui.QPalette.HighlightedText, QtGui.QColor("#ffffff"))
    pal.setColor(QtGui.QPalette.ToolTipBase, QtGui.QColor("#121622"))
    pal.setColor(QtGui.QPalette.ToolTipText, QtGui.QColor("#e6e6e6"))
    pal.setColor(QtGui.QPalette.PlaceholderText, QtGui.QColor("#aeb8cf"))
    app.setPalette(pal)

    # QSS mirrors the launcher look: rounded cards, readable menus, stable contrast.
    app.setStyleSheet("""
        QMainWindow { background: #0f1115; }
        QWidget { color: #e6e6e6; }
        QLabel { color: #e6e6e6; }
        QToolTip { background: #121622; color: #e6e6e6; border: 1px solid #2a3140; padding: 6px; border-radius: 8px; }

        QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox, QDoubleSpinBox, QDateTimeEdit, QComboBox {
            background: #161a22; color: #e6e6e6;
            border: 1px solid #2a3140; border-radius: 8px;
            padding: 6px 8px;
            selection-background-color: #2b64ff;
        }
        QComboBox::drop-down { border: none; width: 28px; }
        QComboBox QAbstractItemView {
            background: #121622; color: #e6e6e6;
            border: 1px solid #2a3140; border-radius: 10px;
            selection-background-color: #2b64ff;
        }

        QPushButton {
            background: #2b64ff; color: white;
            border: none; border-radius: 10px;
            padding: 8px 12px;
        }
        QPushButton:hover { background: #3b74ff; }
        QPushButton:pressed { background: #2456dc; }
        QPushButton:disabled { background: #2a3140; color: #8b93a5; }

        QDialog { background: #0f1115; }
        QGroupBox {
            border: 1px solid #2a3140;
            border-radius: 12px;
            margin-top: 10px;
            padding: 10px;
        }
        QGroupBox::title { subcontrol-origin: margin; left: 12px; padding: 0 6px; color: #cfd6e6; }

        QTabWidget::pane { border: 1px solid #2a3140; border-radius: 12px; background: #121622; }
        QTabBar::tab {
            background: #161a22; color: #cfd6e6;
            border: 1px solid #2a3140;
            border-bottom: none;
            padding: 8px 12px;
            border-top-left-radius: 10px;
            border-top-right-radius: 10px;
            margin-right: 6px;
        }
        QTabBar::tab:selected { background: #121622; color: #ffffff; }

        QTableWidget, QTreeWidget, QListWidget {
            background: #161a22; color: #e6e6e6;
            border: 1px solid #2a3140; border-radius: 10px;
            gridline-color: #2a3140;
            selection-background-color: #2b64ff;
        }
        QHeaderView::section {
            background: #121622; color: #cfd6e6;
            border: none; padding: 8px;
        }

        QMenuBar {
            background: #0f1115;
            color: #e6e6e6;
        }
        QMenuBar::item { padding: 6px 10px; background: transparent; }
        QMenuBar::item:selected { background: #161a22; border-radius: 8px; }
        QMenu {
            background: #121622; color: #e6e6e6;
            border: 1px solid #2a3140; border-radius: 10px;
            padding: 6px;
        }
        QMenu::item { padding: 6px 18px; border-radius: 8px; }
        QMenu::item:selected { background: #2b64ff; color: #ffffff; }

        QScrollBar:vertical {
            background: #121622; width: 12px; margin: 0px; border-radius: 6px;
        }
        QScrollBar::handle:vertical {
            background: #2a3140; min-height: 24px; border-radius: 6px;
        }
        QScrollBar::handle:vertical:hover { background: #3b465b; }
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
        QScrollBar:horizontal {
            background: #121622; height: 12px; margin: 0px; border-radius: 6px;
        }
        QScrollBar::handle:horizontal {
            background: #2a3140; min-width: 24px; border-radius: 6px;
        }
        QScrollBar::handle:horizontal:hover { background: #3b465b; }
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
    """)


# =========================
# App
# =========================


def main():
    log_event("START", cwd=os.getcwd(), argv=sys.argv, python=sys.version)
    log_event("LOG_PATH", path=LOG_PATH)

    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--version", action="store_true")
    parser.add_argument("--mode", default="offline")
    parser.add_argument("--server_url", default="")
    parser.add_argument("--game_id", default="")
    parser.add_argument("--access_token", default="")
    parser.add_argument("--activation_key", default="")
    parser.add_argument("--username", default="local")
    args, _ = parser.parse_known_args()

    if getattr(args, "version", False):
        # Для дебага/CI: py editor.py --version
        print(get_app_version())
        return

    # Обновляем/создаём version.txt рядом с editor.exe (или editor.py).
    write_version_txt()

    ctx = OnlineContext(
        mode=str(args.mode or "offline").lower(),
        server_url=str(args.server_url or ""),
        game_id=str(args.game_id or ""),
        access_token=str(args.access_token or ""),
        activation_key=str(args.activation_key or ""),
        username=str(args.username or "local"),
    )

    app = QtWidgets.QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    if os.path.exists(ICON_PATH):
        try:
            app.setWindowIcon(QtGui.QIcon(ICON_PATH))
        except Exception:
            pass
    apply_dark_theme(app)
    w = MainWindow(online=ctx)
    if os.path.exists(ICON_PATH):
        try:
            w.setWindowIcon(QtGui.QIcon(ICON_PATH))
        except Exception:
            pass
    w.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
